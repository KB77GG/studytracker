#!/usr/bin/env python3
"""
Generate a branded IELTS Key Vocabulary Word document for 睿然国际教育.
Covers all four sections, all band levels (5-6 / 6-7 / 7+).
Organised by thematic topic with exam-style presentation.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, textwrap

# ── brand palette ──
TEAL       = RGBColor(0x2F, 0x8E, 0x87)
TEAL_LIGHT = RGBColor(0xE8, 0xF5, 0xF3)
TEAL_DARK  = RGBColor(0x1A, 0x5C, 0x57)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x2D, 0x2D, 0x2D)
GRAY_TEXT   = RGBColor(0x66, 0x66, 0x66)
ORANGE      = RGBColor(0xE8, 0x7C, 0x2A)
BAND56_BG   = "FFF8E1"   # warm yellow
BAND67_BG   = "E8F5F3"   # light teal
BAND7P_BG   = "EDE7F6"   # light purple

OUT_DIR  = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
OUT_PATH = os.path.join(OUT_DIR, "睿然国际教育_雅思核心词汇手册.docx")

# ────────────────────────────────────────────
# VOCABULARY DATA  (topic → list of entries)
# Each entry: (word, pos, chinese, example, sections, band)
# sections: L=Listening R=Reading W=Writing S=Speaking
# band: "5-6" / "6-7" / "7+"
# ────────────────────────────────────────────

VOCAB = {
    "教育与学习 Education & Learning": [
        ("curriculum", "n.", "课程体系", "The school revised its curriculum to include more practical subjects.", "R W", "5-6"),
        ("academic", "adj.", "学术的", "She has an impressive academic record.", "R W S", "5-6"),
        ("assignment", "n.", "作业；任务", "Students must submit the assignment by Friday.", "L R W", "5-6"),
        ("enrol / enroll", "v.", "注册；入学", "He decided to enrol in a postgraduate programme.", "L R", "5-6"),
        ("semester", "n.", "学期", "The course runs for two semesters.", "L R", "5-6"),
        ("tuition", "n.", "学费；教学", "Tuition fees have risen sharply.", "L R W", "5-6"),
        ("discipline", "n.", "学科；纪律", "History is a discipline that requires critical thinking.", "R W S", "6-7"),
        ("pedagogy", "n.", "教学法", "Modern pedagogy emphasises student-centred learning.", "R W", "7+"),
        ("cognitive", "adj.", "认知的", "Cognitive development is crucial in early childhood.", "R W", "6-7"),
        ("rote learning", "n.", "死记硬背", "Rote learning is less effective than deep understanding.", "R W S", "6-7"),
        ("vocational", "adj.", "职业的", "Vocational training prepares students for specific careers.", "R W S", "6-7"),
        ("literacy", "n.", "读写能力", "Improving literacy rates is a national priority.", "R W", "6-7"),
        ("syllabus", "n.", "教学大纲", "The syllabus covers both theory and practice.", "L R", "5-6"),
        ("autonomous", "adj.", "自主的", "Autonomous learners tend to achieve higher results.", "R W", "7+"),
        ("plagiarism", "n.", "抄袭；剽窃", "Universities have strict policies against plagiarism.", "R W S", "7+"),
    ],

    "科技与创新 Technology & Innovation": [
        ("innovation", "n.", "创新", "Technological innovation drives economic growth.", "R W S", "5-6"),
        ("artificial intelligence", "n.", "人工智能", "Artificial intelligence is transforming many industries.", "R W S", "6-7"),
        ("automation", "n.", "自动化", "Automation has replaced many manual jobs.", "R W S", "6-7"),
        ("digital", "adj.", "数字的", "The digital revolution has changed how we communicate.", "L R W S", "5-6"),
        ("gadget", "n.", "小型电子设备", "Modern gadgets make everyday life more convenient.", "L S", "5-6"),
        ("breakthrough", "n.", "突破", "Scientists announced a major breakthrough in cancer research.", "R W", "6-7"),
        ("obsolete", "adj.", "过时的", "New technology has made typewriters obsolete.", "R W", "7+"),
        ("cutting-edge", "adj.", "尖端的；前沿的", "The lab uses cutting-edge equipment.", "R W S", "6-7"),
        ("algorithm", "n.", "算法", "Social media algorithms determine what content users see.", "R W", "7+"),
        ("cybersecurity", "n.", "网络安全", "Cybersecurity threats are increasing worldwide.", "R W", "7+"),
        ("broadband", "n.", "宽带", "Rural areas often lack reliable broadband access.", "L R", "5-6"),
        ("download", "v./n.", "下载", "You can download the app for free.", "L S", "5-6"),
        ("virtual reality", "n.", "虚拟现实", "Virtual reality is used in medical training.", "R W S", "6-7"),
        ("data privacy", "n.", "数据隐私", "Data privacy has become a major concern.", "R W S", "6-7"),
        ("renewable energy", "n.", "可再生能源", "Renewable energy sources include solar and wind power.", "R W S", "6-7"),
    ],

    "环境与自然 Environment & Nature": [
        ("pollution", "n.", "污染", "Air pollution is a serious problem in big cities.", "R W S", "5-6"),
        ("ecosystem", "n.", "生态系统", "The oil spill damaged the local ecosystem.", "R W", "6-7"),
        ("carbon footprint", "n.", "碳足迹", "We should all try to reduce our carbon footprint.", "R W S", "6-7"),
        ("sustainable", "adj.", "可持续的", "Sustainable development balances growth and environment.", "R W S", "6-7"),
        ("deforestation", "n.", "森林砍伐", "Deforestation contributes to climate change.", "R W", "6-7"),
        ("biodiversity", "n.", "生物多样性", "Biodiversity loss threatens the stability of ecosystems.", "R W", "7+"),
        ("conservation", "n.", "保护；保育", "Wildlife conservation efforts have saved several species.", "R W S", "6-7"),
        ("emission", "n.", "排放", "The government aims to cut carbon emissions by 40%.", "R W", "6-7"),
        ("drought", "n.", "干旱", "The region is suffering from severe drought.", "L R W", "5-6"),
        ("habitat", "n.", "栖息地", "Urban expansion is destroying natural habitats.", "R W", "6-7"),
        ("endangered", "adj.", "濒危的", "The panda is an endangered species.", "R W S", "5-6"),
        ("fossil fuel", "n.", "化石燃料", "Burning fossil fuels releases greenhouse gases.", "R W", "6-7"),
        ("recycle", "v.", "回收利用", "Most households now recycle plastic and paper.", "L S", "5-6"),
        ("ecological", "adj.", "生态的", "The dam caused significant ecological damage.", "R W", "7+"),
        ("carbon-neutral", "adj.", "碳中和的", "The company pledged to become carbon-neutral by 2030.", "R W", "7+"),
    ],

    "社会与文化 Society & Culture": [
        ("community", "n.", "社区；群体", "A strong community supports its members.", "L R W S", "5-6"),
        ("globalisation", "n.", "全球化", "Globalisation has both benefits and drawbacks.", "R W S", "6-7"),
        ("immigration", "n.", "移民", "Immigration policies vary widely between countries.", "R W S", "6-7"),
        ("discrimination", "n.", "歧视", "Discrimination based on gender is illegal.", "R W S", "6-7"),
        ("tradition", "n.", "传统", "Many traditions are passed down through generations.", "L R W S", "5-6"),
        ("inequality", "n.", "不平等", "Income inequality has widened in recent decades.", "R W", "6-7"),
        ("multicultural", "adj.", "多元文化的", "London is one of the most multicultural cities.", "R W S", "6-7"),
        ("stereotype", "n.", "刻板印象", "Media can reinforce harmful stereotypes.", "R W S", "7+"),
        ("assimilate", "v.", "融入；同化", "Immigrants often face pressure to assimilate.", "R W", "7+"),
        ("heritage", "n.", "遗产；传承", "Protecting cultural heritage is essential.", "R W S", "6-7"),
        ("demographic", "adj./n.", "人口统计的", "Demographic changes affect economic planning.", "R W", "7+"),
        ("urbanisation", "n.", "城市化", "Rapid urbanisation creates housing challenges.", "R W", "6-7"),
        ("volunteer", "n./v.", "志愿者；自愿", "She volunteers at a local charity every weekend.", "L R S", "5-6"),
        ("philanthropy", "n.", "慈善事业", "Corporate philanthropy can benefit society.", "R W", "7+"),
        ("social cohesion", "n.", "社会凝聚力", "Festivals promote social cohesion.", "R W", "7+"),
    ],

    "健康与生活方式 Health & Lifestyle": [
        ("obesity", "n.", "肥胖症", "Childhood obesity has become a global concern.", "R W S", "6-7"),
        ("nutrition", "n.", "营养", "Good nutrition is essential for healthy growth.", "L R W S", "5-6"),
        ("sedentary", "adj.", "久坐的", "A sedentary lifestyle increases health risks.", "R W S", "6-7"),
        ("well-being", "n.", "幸福；健康", "Exercise improves both physical and mental well-being.", "R W S", "6-7"),
        ("symptom", "n.", "症状", "Common symptoms include fever and coughing.", "L R", "5-6"),
        ("chronic", "adj.", "慢性的", "Chronic diseases require long-term management.", "R W", "6-7"),
        ("epidemic", "n.", "流行病", "The flu epidemic affected thousands of people.", "R W", "6-7"),
        ("immune system", "n.", "免疫系统", "A balanced diet strengthens the immune system.", "R W", "6-7"),
        ("mental health", "n.", "心理健康", "Mental health awareness has increased significantly.", "R W S", "6-7"),
        ("allergy", "n.", "过敏", "Nut allergies can be life-threatening.", "L R S", "5-6"),
        ("diagnosis", "n.", "诊断", "Early diagnosis improves treatment outcomes.", "R W", "6-7"),
        ("remedy", "n.", "治疗方法；补救", "Herbal remedies are popular in traditional medicine.", "R W S", "6-7"),
        ("preventive", "adj.", "预防性的", "Preventive measures can reduce disease spread.", "R W", "7+"),
        ("rehabilitation", "n.", "康复", "The rehabilitation programme helps patients recover.", "R W", "7+"),
        ("longevity", "n.", "长寿", "Advances in medicine have increased human longevity.", "R W", "7+"),
    ],

    "工作与经济 Work & Economy": [
        ("employment", "n.", "就业", "The government aims to boost employment.", "R W S", "5-6"),
        ("salary", "n.", "薪水", "Her salary increased after the promotion.", "L R S", "5-6"),
        ("recession", "n.", "经济衰退", "The country experienced a severe recession.", "R W", "6-7"),
        ("entrepreneur", "n.", "企业家", "Many entrepreneurs start businesses from home.", "R W S", "6-7"),
        ("inflation", "n.", "通货膨胀", "Inflation has driven up the cost of living.", "R W", "6-7"),
        ("freelance", "adj./v.", "自由职业的", "She works freelance as a graphic designer.", "L R S", "6-7"),
        ("redundancy", "n.", "裁员；冗余", "Hundreds of workers faced redundancy.", "R W", "7+"),
        ("outsource", "v.", "外包", "Many companies outsource their IT services.", "R W", "7+"),
        ("productivity", "n.", "生产力", "New software has improved our productivity.", "R W S", "6-7"),
        ("subsidy", "n.", "补贴", "The government provides subsidies for farmers.", "R W", "7+"),
        ("commodity", "n.", "商品；大宗商品", "Oil is the most traded global commodity.", "R W", "7+"),
        ("workforce", "n.", "劳动力", "The workforce is becoming increasingly diverse.", "R W S", "6-7"),
        ("minimum wage", "n.", "最低工资", "The minimum wage was raised to £12 per hour.", "R W S", "5-6"),
        ("trade deficit", "n.", "贸易逆差", "The trade deficit widened this quarter.", "R W", "7+"),
        ("gig economy", "n.", "零工经济", "The gig economy offers flexibility but less security.", "R W S", "7+"),
    ],

    "城市与交通 Cities & Transport": [
        ("infrastructure", "n.", "基础设施", "Improving infrastructure is key to development.", "R W", "6-7"),
        ("congestion", "n.", "拥堵", "Traffic congestion is a major issue in cities.", "R W S", "6-7"),
        ("commute", "v./n.", "通勤", "She commutes two hours to work every day.", "L R S", "5-6"),
        ("pedestrian", "n.", "行人", "A new pedestrian zone was built in the city centre.", "L R W", "5-6"),
        ("suburb", "n.", "郊区", "Many families prefer living in the suburbs.", "L R S", "5-6"),
        ("public transport", "n.", "公共交通", "Investing in public transport reduces emissions.", "R W S", "5-6"),
        ("urban sprawl", "n.", "城市扩张", "Urban sprawl consumes agricultural land.", "R W", "7+"),
        ("zoning", "n.", "分区规划", "Zoning laws separate residential and industrial areas.", "R W", "7+"),
        ("toll", "n.", "通行费", "Drivers pay a toll to use the highway.", "L R", "6-7"),
        ("sustainable transport", "n.", "可持续交通", "Cycling is a form of sustainable transport.", "R W S", "6-7"),
        ("metro / subway", "n.", "地铁", "The metro system carries millions of passengers daily.", "L R S", "5-6"),
        ("emission-free", "adj.", "零排放的", "The city plans to have an emission-free bus fleet.", "R W", "7+"),
        ("intersection", "n.", "十字路口", "There is a traffic light at the intersection.", "L R", "5-6"),
        ("bypass", "n./v.", "绕行道；绕过", "A new bypass will divert traffic from the town centre.", "R W", "6-7"),
        ("car-sharing", "n.", "共享汽车", "Car-sharing services reduce the number of vehicles.", "R W S", "6-7"),
    ],

    "媒体与传播 Media & Communication": [
        ("broadcast", "v./n.", "广播；播出", "The match was broadcast live on television.", "L R W", "5-6"),
        ("journalism", "n.", "新闻业", "Investigative journalism uncovers important truths.", "R W S", "6-7"),
        ("censorship", "n.", "审查制度", "Censorship restricts freedom of expression.", "R W S", "7+"),
        ("social media", "n.", "社交媒体", "Social media has transformed how people interact.", "R W S", "5-6"),
        ("misinformation", "n.", "错误信息", "Misinformation spreads quickly online.", "R W S", "6-7"),
        ("mainstream", "adj.", "主流的", "Mainstream media sometimes overlooks local issues.", "R W S", "6-7"),
        ("viral", "adj.", "病毒式传播的", "The video went viral within hours.", "R W S", "6-7"),
        ("propaganda", "n.", "宣传；鼓吹", "Propaganda was widely used during the war.", "R W", "7+"),
        ("circulation", "n.", "发行量；流通", "The newspaper has a daily circulation of 500,000.", "R W", "7+"),
        ("tabloid", "n.", "小报", "Tabloids tend to focus on celebrity gossip.", "R W S", "7+"),
        ("advertisement", "n.", "广告", "The advertisement attracted thousands of viewers.", "L R W S", "5-6"),
        ("subscribe", "v.", "订阅", "You can subscribe to the newsletter for free.", "L R", "5-6"),
        ("bias", "n.", "偏见；偏向", "Media bias can influence public opinion.", "R W S", "6-7"),
        ("influencer", "n.", "网红；意见领袖", "Social media influencers shape consumer behaviour.", "R W S", "6-7"),
        ("press freedom", "n.", "新闻自由", "Press freedom is fundamental to democracy.", "R W", "7+"),
    ],

    "犯罪与法律 Crime & Law": [
        ("crime rate", "n.", "犯罪率", "The crime rate has dropped significantly.", "R W S", "5-6"),
        ("offender", "n.", "罪犯", "Young offenders may be sent to reform schools.", "R W S", "6-7"),
        ("penalty", "n.", "处罚；刑罚", "The penalty for speeding can be severe.", "R W S", "6-7"),
        ("rehabilitation", "n.", "改造；康复", "Prison rehabilitation programmes reduce reoffending.", "R W", "6-7"),
        ("deterrent", "n./adj.", "威慑", "Harsh sentences serve as a deterrent.", "R W", "7+"),
        ("legislation", "n.", "立法", "New legislation was introduced to combat fraud.", "R W", "7+"),
        ("verdict", "n.", "裁决", "The jury reached a unanimous verdict.", "R W", "7+"),
        ("surveillance", "n.", "监控", "CCTV surveillance helps prevent crime.", "R W S", "6-7"),
        ("juvenile", "adj./n.", "青少年的", "Juvenile crime is a growing concern.", "R W", "6-7"),
        ("prosecution", "n.", "起诉", "The prosecution presented strong evidence.", "R W", "7+"),
        ("victim", "n.", "受害者", "Support services are available for victims of crime.", "L R W S", "5-6"),
        ("vandalism", "n.", "故意破坏", "Vandalism costs the city millions each year.", "R W S", "6-7"),
        ("fraud", "n.", "欺诈", "Online fraud has increased sharply.", "R W", "6-7"),
        ("witness", "n./v.", "目击者；目睹", "The witness described the suspect to the police.", "L R", "5-6"),
        ("parole", "n.", "假释", "He was released on parole after five years.", "R W", "7+"),
    ],

    "雅思高频学术词汇 Academic & General High-Frequency": [
        ("significant", "adj.", "重要的；显著的", "There has been a significant increase in sales.", "R W", "5-6"),
        ("proportion", "n.", "比例", "A large proportion of students passed the exam.", "R W", "6-7"),
        ("fluctuate", "v.", "波动", "Prices tend to fluctuate throughout the year.", "R W", "6-7"),
        ("tendency", "n.", "趋势；倾向", "There is a tendency for prices to rise in summer.", "R W", "6-7"),
        ("approximately", "adv.", "大约", "Approximately 70% of participants agreed.", "R W", "5-6"),
        ("considerably", "adv.", "相当大地", "Output has increased considerably.", "R W", "6-7"),
        ("nevertheless", "adv.", "然而；尽管如此", "The plan was risky; nevertheless, they proceeded.", "R W", "7+"),
        ("whereas", "conj.", "然而；鉴于", "Urban areas are crowded, whereas rural areas are quiet.", "R W", "6-7"),
        ("subsequently", "adv.", "随后", "He graduated and subsequently found a job.", "R W", "7+"),
        ("phenomenon", "n.", "现象", "Climate change is a complex phenomenon.", "R W", "6-7"),
        ("perspective", "n.", "视角；观点", "We should consider the issue from multiple perspectives.", "R W S", "6-7"),
        ("hypothesis", "n.", "假设", "The hypothesis was tested through experiments.", "R W", "7+"),
        ("correlation", "n.", "相关性", "There is a strong correlation between exercise and health.", "R W", "7+"),
        ("criteria", "n.", "标准（复数）", "Applicants must meet all the criteria.", "R W", "6-7"),
        ("implication", "n.", "含义；影响", "The findings have important implications for policy.", "R W", "7+"),
        ("inevitably", "adv.", "不可避免地", "Costs will inevitably rise over time.", "R W", "7+"),
        ("underlying", "adj.", "根本的；潜在的", "The underlying cause of the problem is poverty.", "R W", "7+"),
        ("implement", "v.", "实施", "The government plans to implement new regulations.", "R W", "6-7"),
        ("adequate", "adj.", "足够的；充分的", "The school lacks adequate funding.", "R W", "6-7"),
        ("controversial", "adj.", "有争议的", "The decision was highly controversial.", "R W S", "6-7"),
    ],
}


# ────────────────────────────────────────────
# Helpers (reuse branding from TA doc)
# ────────────────────────────────────────────

def _set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    xml = f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}"'
    if line:
        xml += f' w:line="{line}" w:lineRule="auto"'
    xml += '/>'
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    pPr.append(parse_xml(xml))

def _add_hr(doc, color="2F8E87", sz=6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'))
    _set_spacing(p, 0, 100)
    return p

def _run(para, text, size=10.5, bold=False, color=DARK_TEXT, italic=False, font="微软雅黑"):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if italic:
        r.font.italic = True
    return r

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        + ''.join(f'<w:{e} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                  for e in ("top","left","bottom","right","insideH","insideV"))
        + '</w:tblBorders>')
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(bdr)


def _set_table_borders(table, header_color="2F8E87", inner_color="BDDBD8"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{header_color}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="{inner_color}"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="{inner_color}"/>'
        f'</w:tblBorders>')
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(bdr)


# ────────────────────────────────────────────
# Section heading with badge (same as TA doc)
# ────────────────────────────────────────────

def add_topic_heading(doc, number, title_cn, title_en=""):
    spacer = doc.add_paragraph()
    _set_spacing(spacer, 140, 0)

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(1.5)
    t.columns[1].width = Cm(14.5)

    c0 = t.rows[0].cells[0]; c0.text = ""
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p0, str(number).zfill(2), 13, True, WHITE, font="Arial")
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(c0, "2F8E87")

    c1 = t.rows[0].cells[1]; c1.text = ""
    p1 = c1.paragraphs[0]
    _run(p1, f"  {title_cn}", 15, True, TEAL_DARK)
    if title_en:
        _run(p1, f"  {title_en}", 11, False, GRAY_TEXT, font="Arial")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(c1, "E8F5F3")
    _remove_table_borders(t)
    _add_hr(doc, "2F8E87", 4)


# ────────────────────────────────────────────
# Build the vocab table for one topic
# ────────────────────────────────────────────

BAND_COLORS = {"5-6": BAND56_BG, "6-7": BAND67_BG, "7+": BAND7P_BG}
BAND_LABEL  = {"5-6": "Band 5-6  基础", "6-7": "Band 6-7  进阶", "7+": "Band 7+  高阶"}
BAND_TEXT_COLOR = {
    "5-6": RGBColor(0xE6, 0x8A, 0x00),
    "6-7": TEAL,
    "7+":  RGBColor(0x7B, 0x1F, 0xA2),
}

def build_vocab_table(doc, entries):
    """Create a 6-column vocab table: # | Word | POS | Chinese | Example | Sections"""
    headers = ["#", "Word / Phrase", "词性", "中文释义", "真题例句", "考试板块"]
    col_w   = [0.7,   3.2,           0.8,    2.3,       6.5,       2.5]

    # Sort by band then word
    band_order = {"5-6": 0, "6-7": 1, "7+": 2}
    entries_sorted = sorted(entries, key=lambda e: (band_order.get(e[5], 9), e[0].lower()))

    table = doc.add_table(rows=1 + len(entries_sorted), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, 9, True, WHITE)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(c, "2F8E87")
        c.width = Cm(col_w[i])

    # Data rows
    prev_band = None
    seq = 0
    for idx, (word, pos, cn, example, sections, band) in enumerate(entries_sorted):
        row = table.rows[idx + 1]
        seq += 1

        # Band divider: insert a merged label row? Not easy with python-docx.
        # Instead colour-code rows by band.
        bg = BAND_COLORS.get(band, "FFFFFF")
        band_color = BAND_TEXT_COLOR.get(band, DARK_TEXT)

        vals = [str(seq), word, pos, cn, example, sections]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for ci, val in enumerate(vals):
            c = row.cells[ci]
            c.text = ""
            c.width = Cm(col_w[ci])
            p = c.paragraphs[0]
            p.alignment = aligns[ci]
            _set_spacing(p, 30, 30)

            if ci == 1:  # word column - bold teal
                _run(p, val, 9.5, True, TEAL_DARK, font="Arial")
            elif ci == 4:  # example - italic
                _run(p, val, 8.5, False, GRAY_TEXT, italic=True, font="Arial")
            elif ci == 5:  # sections
                _run(p, val, 8.5, True, band_color, font="Arial")
            else:
                _run(p, val, 9, False, DARK_TEXT)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(c, bg)

    _set_table_borders(table)

    # Band legend below table
    p_legend = doc.add_paragraph()
    p_legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_legend.paragraph_format.space_before = Pt(6)
    p_legend.paragraph_format.space_after = Pt(2)
    _run(p_legend, "  ■", 9, True, BAND_TEXT_COLOR["5-6"], font="Arial")
    _run(p_legend, " Band 5-6 基础   ", 8, False, GRAY_TEXT)
    _run(p_legend, "■", 9, True, BAND_TEXT_COLOR["6-7"], font="Arial")
    _run(p_legend, " Band 6-7 进阶   ", 8, False, GRAY_TEXT)
    _run(p_legend, "■", 9, True, BAND_TEXT_COLOR["7+"], font="Arial")
    _run(p_legend, " Band 7+ 高阶", 8, False, GRAY_TEXT)

    return table


# ────────────────────────────────────────────
# MAIN builder
# ────────────────────────────────────────────

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(29.7)    # A4 landscape
    sec.page_height = Cm(21)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)

    # ── COVER ──
    for _ in range(3):
        s = doc.add_paragraph(); _set_spacing(s, 0, 0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "睿然国际教育", 32, True, TEAL)
    _set_spacing(p, 0, 80)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "RUIRAN INTERNATIONAL EDUCATION", 11, False, GRAY_TEXT, font="Arial")
    _set_spacing(p, 0, 300)

    _add_hr(doc, "2F8E87", 8)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "IELTS 雅思核心词汇手册", 28, True, TEAL_DARK)
    _set_spacing(p, 200, 80)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "真题高频 · 考点精编 · 全科覆盖", 16, False, TEAL)
    _set_spacing(p, 0, 60)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Band 5-6  ·  Band 6-7  ·  Band 7+", 13, True, GRAY_TEXT, font="Arial")
    _set_spacing(p, 0, 60)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Listening  ·  Reading  ·  Writing  ·  Speaking", 12, False, GRAY_TEXT, font="Arial")
    _set_spacing(p, 0, 200)

    _add_hr(doc, "2F8E87", 8)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "内部教学资料 · 请勿外传", 10, False, ORANGE, italic=True)

    doc.add_page_break()

    # ── HOW TO USE ──
    add_topic_heading(doc, 0, "使用说明", "How to Use This Booklet")

    for text in [
        "本手册精选雅思考试真题高频核心词汇，按主题分类，覆盖听说读写四科。",
        "每个词条包含：单词/短语、词性、中文释义、真题风格例句、适用考试板块。",
    ]:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(4)
        pp.paragraph_format.line_spacing = Pt(22)
        _run(pp, text, 10.5, False, DARK_TEXT)

    # band legend box
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(8)
    pPr = pp._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="2F8E87"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5F3" w:val="clear"/>'))
    _run(pp, "词汇按难度标注三个等级，以行背景色区分：\n", 10, False, TEAL_DARK)
    _run(pp, "  ■ Band 5-6 基础词汇（暖黄底色）", 10, True, BAND_TEXT_COLOR["5-6"])
    _run(pp, "   ", 10)
    _run(pp, "■ Band 6-7 进阶词汇（青绿底色）", 10, True, BAND_TEXT_COLOR["6-7"])
    _run(pp, "   ", 10)
    _run(pp, "■ Band 7+ 高阶词汇（淡紫底色）", 10, True, BAND_TEXT_COLOR["7+"])

    pp2 = doc.add_paragraph()
    pp2.paragraph_format.space_before = Pt(8)
    _run(pp2, "考试板块缩写：", 10, True, TEAL_DARK)
    _run(pp2, " L = Listening　R = Reading　W = Writing　S = Speaking", 10, False, GRAY_TEXT, font="Arial")

    doc.add_page_break()

    # ── VOCAB SECTIONS ──
    for idx, (topic, entries) in enumerate(VOCAB.items(), 1):
        # Split topic into CN and EN parts
        parts = topic.split(" ", 1)
        cn_title = parts[0] if len(parts) >= 1 else topic
        en_title = parts[1] if len(parts) >= 2 else ""

        add_topic_heading(doc, idx, cn_title, en_title)
        build_vocab_table(doc, entries)

        # Add page break after every 2nd topic to keep layout clean
        if idx < len(VOCAB):
            doc.add_page_break()

    # ── BACK COVER ──
    doc.add_page_break()
    for _ in range(4):
        s = doc.add_paragraph(); _set_spacing(s, 0, 0)

    _add_hr(doc, "2F8E87", 8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "睿然国际教育", 20, True, TEAL)
    _set_spacing(p, 160, 80)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "用科技赋能教育，让过程被看见，让结果自然发生。", 11, False, GRAY_TEXT, italic=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 120, 0)
    _run(p, "© 睿然国际教育  ·  内部教学资料  ·  版权所有", 9, False, GRAY_TEXT)

    _add_hr(doc, "2F8E87", 8)

    doc.save(OUT_PATH)
    print(f"✅  Saved → {OUT_PATH}")


if __name__ == "__main__":
    build()
