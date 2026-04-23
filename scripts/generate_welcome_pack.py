#!/usr/bin/env python3
"""Generate branded New Student Welcome Pack for 睿然国际教育."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

TEAL=RGBColor(0x2F,0x8E,0x87);TD=RGBColor(0x1A,0x5C,0x57)
W=RGBColor(0xFF,0xFF,0xFF);DK=RGBColor(0x2D,0x2D,0x2D)
GR=RGBColor(0x66,0x66,0x66);OR=RGBColor(0xE8,0x7C,0x2A)

OUT=os.path.join(os.path.dirname(os.path.dirname(__file__)),"docs","睿然国际教育_新生入学指南.docx")

def _sh(c,h): c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}" w:val="clear"/>'))
def _sp(p,b=0,a=0):
    pPr=p._p.get_or_add_pPr()
    for o in pPr.findall(qn('w:spacing')): pPr.remove(o)
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="{b}" w:after="{a}"/>'))
def _hr(d,c="2F8E87",s=6):
    p=d.add_paragraph();p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{s}" w:space="1" w:color="{c}"/></w:pBdr>'));_sp(p,0,100)
def _r(p,t,sz=10.5,b=False,c=DK,i=False,f="微软雅黑"):
    r=p.add_run(t);r.font.size=Pt(sz);r.font.bold=b;r.font.color.rgb=c;r.font.name=f
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if i: r.font.italic=True
    return r
def _nb(t):
    tP=t._tbl.tblPr
    if tP is None: tP=parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    b=parse_xml(f'<w:tblBorders {nsdecls("w")}>'+''.join(f'<w:{e} w:val="none" w:sz="0" w:space="0" w:color="auto"/>' for e in("top","left","bottom","right","insideH","insideV"))+'</w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(b)
def _tb(t):
    tP=t._tbl.tblPr
    if tP is None: tP=parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    b=parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:left w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:right w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:insideH w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/><w:insideV w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/></w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(b)

def hd(d,n,cn,en=""):
    s=d.add_paragraph();_sp(s,140,0)
    t=d.add_table(rows=1,cols=2);t.autofit=False;t.columns[0].width=Cm(1.5);t.columns[1].width=Cm(14.5)
    c0=t.rows[0].cells[0];c0.text="";p0=c0.paragraphs[0];p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p0,str(n).zfill(2),13,True,W,f="Arial");c0.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c0,"2F8E87")
    c1=t.rows[0].cells[1];c1.text="";p1=c1.paragraphs[0]
    _r(p1,f"  {cn}",15,True,TD)
    if en: _r(p1,f"  {en}",11,False,GR,f="Arial")
    c1.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c1,"E8F5F3");_nb(t);_hr(d,"2F8E87",4)

def tbl(d,headers,rows,widths=None):
    t=d.add_table(rows=1+len(rows),cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text="";p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _r(p,h,9.5,True,W);c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c,"2F8E87")
        if widths: c.width=Cm(widths[i])
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci];c.text=""
            if widths: c.width=Cm(widths[ci])
            p=c.paragraphs[0];_sp(p,40,40)
            if ci==0: _r(p,v,9.5,True,TD)
            else: _r(p,v,9.5,False,DK)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            if ri%2==1: _sh(c,"F0FAF9")
    _tb(t);return t

def callout(d,t):
    p=d.add_paragraph();p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.left_indent=Cm(1);pPr=p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="2F8E87"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5F3" w:val="clear"/>'))
    _r(p,t,10,False,TD,i=True)

def body(d,t):
    p=d.add_paragraph();p.paragraph_format.space_before=Pt(4);p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.line_spacing=Pt(22);_r(p,t,10.5,False,DK)

def bullet(d,t):
    p=d.add_paragraph();p.paragraph_format.space_before=Pt(2);p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.line_spacing=Pt(20);p.paragraph_format.left_indent=Cm(1.2)
    p.paragraph_format.first_line_indent=Cm(-0.5)
    # parse **bold**
    parts=t.split("**")
    for i,part in enumerate(parts):
        if not part: continue
        if i%2==1: _r(p,part,10,True,TD)
        else: _r(p,part,10,False,DK)

def build():
    d=Document()
    sec=d.sections[0];sec.page_width=Cm(21);sec.page_height=Cm(29.7)
    sec.top_margin=Cm(2);sec.bottom_margin=Cm(2);sec.left_margin=Cm(2.5);sec.right_margin=Cm(2.5)

    # COVER
    for _ in range(4): s=d.add_paragraph();_sp(s,0,0)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"睿然国际教育",30,True,TEAL);_sp(p,0,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"RUIRAN INTERNATIONAL EDUCATION",11,False,GR,f="Arial");_sp(p,0,300)
    _hr(d,"2F8E87",8)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"新生入学指南",26,True,TD);_sp(p,200,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"Welcome Guide",18,False,TEAL,f="Arial");_sp(p,0,120)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"开启你的备考之旅",14,False,GR,i=True);_sp(p,0,200)
    _hr(d,"2F8E87",8)
    d.add_page_break()

    # S1: Welcome
    hd(d,1,"欢迎来到睿然","Welcome to Ruiran")
    body(d,"欢迎加入睿然国际教育大家庭！从今天起，你不再是一个人在战斗。")
    body(d,"在睿然，我们相信每一个学生都有无限可能。我们的使命是通过个性化教学、数据驱动的学习追踪和专属助教陪伴，帮助你高效备考、科学提分。")
    body(d,"我们与众不同之处在于：你的每一份努力都会被记录，你的每一个薄弱点都会被关注，你的每一步进步都会被看见。这不只是一句口号——我们用技术和制度来保证这一切真实发生。")
    callout(d,"准备好了吗？让我们一起开始这段旅程！")
    d.add_page_break()

    # S2: Learning team
    hd(d,2,"你的学习团队","Your Learning Team")
    body(d,"在睿然，你身边有一支专业的团队为你保驾护航：")
    tbl(d,
        ["角色","职责","你会在哪里找到他们"],
        [
            ["主讲教师","课堂授课、教学规划、阶段性评估","课堂上 / 系统消息"],
            ["专属助教（Coach）","课后任务跟进、每日作业批改、学习督导","微信 / 小程序反馈"],
            ["批改助教（Reviewer）","专项作业精批、正确率评估、错因分析","小程序批改记录"],
            ["排课专员","课程安排、时间协调、调课处理","微信沟通"],
        ],
        widths=[3.5,6,6.5],
    )
    s=d.add_paragraph();_sp(s,60,0)
    callout(d,"你的助教会在入学后第一周主动联系你，请保持微信畅通！")
    d.add_page_break()

    # S3: Mini-program guide
    hd(d,3,"学习追踪小程序使用指南","Mini-Program Guide")

    p=d.add_paragraph();p.paragraph_format.space_before=Pt(10);p.paragraph_format.space_after=Pt(4)
    _r(p,"3.1  如何登录",12,True,TEAL)
    for step in [
        "1.  在微信搜索「睿然学习追踪」小程序",
        "2.  一键微信授权登录",
        '3.  选择「学生」角色',
        "4.  输入你的姓名完成绑定",
    ]: bullet(d, step)

    p=d.add_paragraph();p.paragraph_format.space_before=Pt(14);p.paragraph_format.space_after=Pt(4)
    _r(p,"3.2  每日学习流程",12,True,TEAL)
    for step in [
        "1.  查看今日任务清单（每天由助教根据教学进度制定）",
        "2.  按顺序完成每项任务",
        "3.  提交完成证据（拍照/录音/文字）",
        "4.  等待助教批改与反馈",
        "5.  根据反馈进行订正或复习",
    ]: bullet(d, step)

    p=d.add_paragraph();p.paragraph_format.space_before=Pt(14);p.paragraph_format.space_after=Pt(4)
    _r(p,"3.3  任务状态说明",12,True,TEAL)
    tbl(d,
        ["状态","含义","你需要做什么"],
        [
            ["待完成","任务已布置，等待你开始","开始做任务"],
            ["进行中","你已开始，计时中","继续完成"],
            ["已提交","你已提交，等待批改","等待助教反馈"],
            ["已通过 ✅","批改通过","进入下一项任务"],
            ["部分通过 ⚠️","基本完成，需改进","查看反馈，修正后重新提交"],
            ["未通过 ❌","需要重做","仔细阅读反馈，重新完成"],
        ],
        widths=[3.5,5,7.5],
    )
    d.add_page_break()

    # S4: After-class support
    hd(d,4,"课后辅导安排","After-Class Support")
    body(d,"课后辅导是睿然教学体系的核心环节。以下是你将享受到的课后支持：")
    for item in [
        "✅ **一对一绑定**：你的专属助教会长期跟进你的学习，深入了解你的习惯和薄弱点",
        "✅ **每日批改反馈**：作业提交后当日内收到批改，不用等到下次上课",
        "✅ **多维度反馈**：文字批注、图片标注、语音点评——不只是对错，更有方法指导",
        "✅ **每节课后反馈**：授课教师提交结构化反馈（作业情况、课堂表现、建议），自动推送给家长",
    ]: bullet(d, item)
    s=d.add_paragraph();_sp(s,60,0)
    callout(d,"有任何学习上的问题，第一时间联系你的专属助教！")
    d.add_page_break()

    # S5: Rules
    hd(d,5,"学习规范与期望","Rules & Expectations")
    body(d,"为了确保最佳学习效果，请遵守以下规范：")
    tbl(d,
        ["规范","具体要求"],
        [
            ["出勤要求","按时上课，如需请假请提前一天通知排课专员"],
            ["作业提交","每日任务须在当天 22:00 前提交"],
            ["证据规范","拍照清晰、录音完整、文字认真"],
            ["订正要求", '收到"部分通过"或"未通过"反馈后，24 小时内完成订正'],
            ["沟通礼仪","尊重老师和助教，有问题积极沟通"],
            ["手机管理","上课期间手机静音，专注学习"],
        ],
        widths=[3.5,12.5],
    )
    d.add_page_break()

    # S6: FAQ
    hd(d,6,"常见问题","FAQ")

    faqs = [
        ("Q: 如何查看老师的批改反馈？",
         "A: 打开小程序 → 点击对应任务 → 查看批改详情（文字/图片/语音）。"),
        ("Q: 任务来不及做完怎么办？",
         "A: 先提交已完成的部分，并在留言中说明情况，助教会根据实际情况调整。"),
        ("Q: 我可以提前做明天的任务吗？",
         "A: 可以！在小程序中可以看到未来已布置的任务。"),
        ("Q: 家长可以看到我的学习情况吗？",
         "A: 是的，家长可以通过家长端小程序实时查看你的学习进度和老师反馈。"),
        ("Q: 如何联系我的助教？",
         "A: 入学时助教会添加你的微信，也可以通过小程序消息联系。"),
        ("Q: 遇到不会的题怎么办？",
         "A: 先尝试独立思考，如果仍有困难，在提交时附上你的疑问，助教会在反馈中解答。"),
    ]
    for q, a in faqs:
        callout(d, q)
        body(d, a)
    d.add_page_break()

    # S7: Timeline
    hd(d,7,"你的备考时间线","Study Timeline")
    body(d,"以下是一个典型的 12 周备考参考规划（实际安排以你的课程计划为准）：")
    tbl(d,
        ["阶段","时间","目标","重点"],
        [
            ["基础巩固期","第 1-4 周","夯实基础，养成习惯","词汇积累、语法巩固、每日打卡习惯"],
            ["能力提升期","第 5-8 周","单项突破，查漏补缺","听说读写分项训练、真题练习"],
            ["冲刺强化期","第 9-12 周","模考实战，稳定发挥","全真模考、错题回顾、考试策略"],
            ["考前调整期","考前 1 周","调整状态，自信应考","重点回顾、心态调整、考场技巧"],
        ],
        widths=[3,2.5,4,6.5],
    )
    s=d.add_paragraph();_sp(s,80,0)
    callout(d,"时间线仅供参考，你的助教会根据你的实际情况制定个性化学习计划。")

    # Closing
    d.add_page_break()
    for _ in range(3): s=d.add_paragraph();_sp(s,0,0)
    _hr(d,"2F8E87",8)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_sp(p,160,60)
    _r(p,"致每一位新同学",16,True,TEAL)

    for txt in [
        "每一天的坚持都在为你的目标添砖加瓦。",
        "我们的系统会记录你的每一份努力——让付出被看见，让进步有迹可循。",
        "加油，未来可期！",
    ]:
        p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(6)
        _r(p,txt,11,False,TD,i=True)

    _hr(d,"2F8E87",8)
    s=d.add_paragraph();_sp(s,100,0)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"睿然国际教育",20,True,TEAL);_sp(p,60,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"用科技赋能教育，让过程被看见，让结果自然发生。",11,False,GR,i=True)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_sp(p,120,0)
    _r(p,"© 睿然国际教育  ·  版权所有",9,False,GR)
    _hr(d,"2F8E87",8)

    d.save(OUT);print(f"✅  Saved → {OUT}")

if __name__=="__main__": build()
