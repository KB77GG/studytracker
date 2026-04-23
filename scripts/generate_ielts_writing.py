#!/usr/bin/env python3
"""Generate branded IELTS Writing High-Score Templates document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

TEAL = RGBColor(0x2F, 0x8E, 0x87)
TEAL_LIGHT = RGBColor(0xE8, 0xF5, 0xF3)
TEAL_DARK = RGBColor(0x1A, 0x5C, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)
ORANGE = RGBColor(0xE8, 0x7C, 0x2A)

OUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs",
                   "睿然国际教育_雅思写作高分模板.docx")

# ── helpers (same as other branded docs) ──

def _shading(cell, h):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}" w:val="clear"/>'))

def _sp(p, b=0, a=0):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
    pPr.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="{b}" w:after="{a}"/>'))

def _hr(doc, c="2F8E87", s=6):
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{s}" '
        f'w:space="1" w:color="{c}"/></w:pBdr>'))
    _sp(p, 0, 100)

def _r(p, t, sz=10.5, b=False, c=DARK_TEXT, i=False, f="微软雅黑"):
    r = p.add_run(t); r.font.size=Pt(sz); r.font.bold=b
    r.font.color.rgb=c; r.font.name=f
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i: r.font.italic=True
    return r

def _no_borders(t):
    tP = t._tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}/>');
    bdr = parse_xml(f'<w:tblBorders {nsdecls("w")}>' +
        ''.join(f'<w:{e} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                for e in ("top","left","bottom","right","insideH","insideV")) +
        '</w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(bdr)

def _tbl_borders(t):
    tP = t._tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}/>');
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'<w:insideH w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/>'
        f'<w:insideV w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/>'
        f'</w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(bdr)

def heading(doc, num, cn, en=""):
    s = doc.add_paragraph(); _sp(s, 140, 0)
    t = doc.add_table(rows=1, cols=2); t.autofit=False
    t.columns[0].width=Cm(1.5); t.columns[1].width=Cm(14.5)
    c0=t.rows[0].cells[0]; c0.text=""
    p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p0, str(num).zfill(2), 13, True, WHITE, f="Arial")
    c0.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; _shading(c0,"2F8E87")
    c1=t.rows[0].cells[1]; c1.text=""
    p1=c1.paragraphs[0]
    _r(p1, f"  {cn}", 15, True, TEAL_DARK)
    if en: _r(p1, f"  {en}", 11, False, GRAY_TEXT, f="Arial")
    c1.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; _shading(c1,"E8F5F3")
    _no_borders(t); _hr(doc,"2F8E87",4)

def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _r(p, h, 9.5, True, WHITE); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        _shading(c,"2F8E87")
        if widths: c.width=Cm(widths[i])
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=""
            if widths: c.width=Cm(widths[ci])
            p=c.paragraphs[0]; _sp(p,40,40)
            # bold first column
            if ci==0:
                _r(p, val, 9.5, True, TEAL_DARK)
            elif ci==len(headers)-1 and len(headers)>=4:
                _r(p, val, 9, False, GRAY_TEXT, i=True)
            else:
                _r(p, val, 9.5, False, DARK_TEXT)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            if ri%2==1: _shading(c,"F0FAF9")
    _tbl_borders(t)
    return t

def callout(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.left_indent=Cm(1)
    pPr=p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" '
        f'w:space="8" w:color="2F8E87"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5F3" w:val="clear"/>'))
    _r(p, txt, 10, False, TEAL_DARK, i=True)

def body(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4)
    p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=Pt(22)
    _r(p, txt, 10.5, False, DARK_TEXT)

def sub_heading(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10)
    p.paragraph_format.space_after=Pt(4)
    _r(p, txt, 12, True, TEAL)

# ── BUILD ──

def build():
    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # COVER
    for _ in range(3): s=doc.add_paragraph(); _sp(s,0,0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"睿然国际教育",30,True,TEAL); _sp(p,0,80)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"RUIRAN INTERNATIONAL EDUCATION",11,False,GRAY_TEXT,f="Arial"); _sp(p,0,300)
    _hr(doc,"2F8E87",8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"IELTS 雅思写作高分模板",26,True,TEAL_DARK); _sp(p,200,80)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"Task 1 图表描述 · Task 2 议论文",16,False,TEAL); _sp(p,0,60)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"Band 6.0 - 8.0+ 适用",13,True,GRAY_TEXT,f="Arial"); _sp(p,0,200)
    _hr(doc,"2F8E87",8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"内部教学资料 · 请勿外传",10,False,ORANGE,i=True)
    doc.add_page_break()

    # USAGE GUIDE
    heading(doc, 0, "使用说明", "How to Use")
    body(doc, "本手册汇集雅思写作 Task 1 与 Task 2 各类高分句型模板，按功能分类整理。")
    body(doc, "每个模板均标注适用场景，方括号 [ ] 内的内容需根据实际题目替换。建议先理解模板逻辑，再通过真题反复练习内化。")
    callout(doc, "提示：模板是工具而非万能钥匙。考官看重的是灵活运用与逻辑清晰，切忌死记硬背。")
    doc.add_page_break()

    # ══════ TASK 1 ══════

    # S1: Intro paraphrasing
    heading(doc, 1, "开头段改写模板", "Task 1 · Introduction Paraphrasing")
    body(doc, "Task 1 开头段的核心任务是改写题目描述，避免原文照抄。以下模板覆盖各类图表：")
    make_table(doc,
        ["图表类型", "模板句型", "关键替换词"],
        [
            ["折线图", "The line graph illustrates [变化对象] over the period from [起始] to [结束].",
             "illustrates → depicts / shows / presents"],
            ["柱状图", "The bar chart compares [比较对象] in terms of [指标] across [类别/时间].",
             "compares → contrasts / draws a comparison of"],
            ["饼图", "The pie chart(s) show(s) the proportion of [对象] in [时间/地点].",
             "proportion → percentage / share / breakdown"],
            ["表格", "The table provides data on [内容] in [范围] between [时间1] and [时间2].",
             "provides → presents / gives information about"],
            ["地图", "The two maps illustrate the changes in [地点] between [时间1] and [时间2].",
             "changes → developments / transformations"],
            ["流程图", "The diagram illustrates the process by which [过程名称].",
             "illustrates → outlines / shows the stages of"],
            ["混合图", "The [图表1] and [图表2] together provide information about [主题].",
             "provide → present / give an overview of"],
            ["通用改写", "shows → illustrates; about → regarding; number → figure; percentage → proportion",
             "给出的 → 提供的；不同的 → 各种各样的"],
        ],
        widths=[2.5, 8, 5.5],
    )
    doc.add_page_break()

    # S2: Trends
    heading(doc, 2, "趋势描述句型", "Task 1 · Trend Description")
    body(doc, "描述数据变化是 Task 1 的核心能力。掌握以下句型可以灵活表达各种趋势：")
    make_table(doc,
        ["趋势类型", "句型模板", "同义替换 / 程度修饰"],
        [
            ["上升", "[主语] increased / rose / grew from [数值] to [数值].",
             "surged / climbed / soared（急剧）; edged up（微升）"],
            ["下降", "[主语] decreased / fell / dropped to [数值] by [年份].",
             "plummeted / plunged（暴跌）; dipped / declined（缓降）"],
            ["急剧变化", "[主语] experienced a dramatic/sharp/significant rise/fall.",
             "dramatic → substantial / remarkable / steep"],
            ["缓慢变化", "[主语] saw a gradual/steady/moderate increase/decrease.",
             "gradual → slow / marginal / slight"],
            ["保持稳定", "[主语] remained stable/constant/unchanged at [数值].",
             "levelled off / plateaued / stayed at around"],
            ["波动", "[主语] fluctuated between [数值] and [数值] during this period.",
             "fluctuated → varied / oscillated / was erratic"],
            ["达到峰值", "[主语] reached a peak / peaked at [数值] in [时间].",
             "hit the highest point / climbed to a high of"],
            ["降至最低", "[主语] fell to its lowest point / bottomed out at [数值].",
             "reached a trough / hit a low of / sank to"],
            ["翻倍", "[主语] doubled / tripled over the period.",
             "increased twofold / threefold / grew by 200%"],
            ["从…到…", "There was a [adj.] [noun] in [主语], from [数值] to [数值].",
             "名词化：rise, increase, growth, decline, drop, fall"],
            ["时间短语", "Over the period / Between ... and ... / During the decade / By [年份]",
             "throughout / across / spanning / in the space of"],
            ["复合趋势", "After rising to [数值], [主语] then fell sharply before stabilising at [数值].",
             "having + p.p. 结构 / followed by / prior to"],
        ],
        widths=[2.5, 7, 6.5],
    )
    doc.add_page_break()

    # S3: Data comparison
    heading(doc, 3, "数据描述与比较", "Task 1 · Data Comparison")
    body(doc, "精确的数据比较和引用能力是 Task 1 拿高分的关键：")
    make_table(doc,
        ["功能", "句型模板", "使用场景"],
        [
            ["最高值", "[A] had the highest [指标], at [数值], followed by [B] at [数值].",
             "比较多组数据中的最大值"],
            ["最低值", "[A] recorded the lowest figure, standing at only [数值].",
             "突出最小值"],
            ["倍数", "[A] was approximately twice / three times as high as [B].",
             "两组数据之间的倍数关系"],
            ["差距", "There was a significant gap between [A] ([数值]) and [B] ([数值]).",
             "强调数据差异"],
            ["相近", "[A] and [B] had almost identical figures, at around [数值].",
             "两组数据接近时"],
            ["占比", "[A] accounted for / made up / represented [比例] of the total.",
             "饼图、百分比描述"],
            ["约数", "Approximately / roughly / just over / just under [数值]",
             "非精确数据引用"],
            ["比较级", "[A] was considerably / slightly / marginally higher than [B].",
             "两组数据直接比较"],
            ["排序", "[A] ranked first, followed by [B], while [C] came last.",
             "多组数据排名"],
            ["总量", "The total amount of [主语] reached [数值] by [时间].",
             "汇总数据"],
        ],
        widths=[2.5, 7.5, 6],
    )
    doc.add_page_break()

    # S4: Overview
    heading(doc, 4, "概述段模板", "Task 1 · Overview")
    callout(doc, "概述段（Overview）是 Task 1 评分的关键要素。没有概述段很难拿到 Band 6 以上。")
    make_table(doc,
        ["模板类型", "句型模板"],
        [
            ["总体趋势", "Overall, it is clear that [主要趋势1], while [主要趋势2]."],
            ["最显著特征", "The most striking feature is that [特征描述]."],
            ["对比总结", "In general, [A] was significantly higher than [B] throughout the period."],
            ["变化方向", "Overall, [主语] showed an upward / downward trend over the period."],
            ["类别差异", "It is noticeable that [类别A] dominated, whereas [类别B] remained relatively low."],
            ["地图概述", "Overall, the area underwent significant development / transformation between [时间1] and [时间2]."],
            ["流程概述", "Overall, the process involves [数量] main stages, beginning with [起点] and ending with [终点]."],
            ["通用框架", "From an overall perspective, two key trends can be identified: firstly, ... ; secondly, ... ."],
        ],
        widths=[3.5, 12.5],
    )
    doc.add_page_break()

    # ══════ TASK 2 ══════

    # S5: Task 2 Intro
    heading(doc, 5, "开头段模板", "Task 2 · Introduction")
    body(doc, "Task 2 开头段需完成两件事：①改写题目背景 ②表明自己的立场/文章结构。")
    make_table(doc,
        ["题型", "开头段模板", "关键要素"],
        [
            ["同意/不同意",
             "It is often argued that [观点改写]. While I understand this perspective, I [agree/disagree] that [你的立场].",
             "明确表态 + 改写题目"],
            ["讨论双方观点",
             "There is an ongoing debate about whether [话题]. This essay will discuss both viewpoints before presenting my own opinion.",
             "预告结构 + 改写"],
            ["优缺点",
             "[话题] has become increasingly common in recent years. This essay will examine both the advantages and disadvantages of this trend.",
             "背景引入 + 预告"],
            ["问题/解决方案",
             "[问题描述] has become a pressing issue in modern society. This essay will explore the main causes and propose possible solutions.",
             "问题陈述 + 预告"],
            ["双问题",
             "[话题] is a topic that generates considerable discussion. This essay will address [问题1] and then consider [问题2].",
             "话题引入 + 双问题预告"],
            ["报告类",
             "In many countries, [现象]. This essay will analyse the reasons for this trend and evaluate its effects.",
             "现象描述 + 分析预告"],
            ["积极/消极发展",
             "[现象] has become a defining feature of the modern age. In my view, this is largely a [positive/negative] development.",
             "明确态度"],
            ["通用万能开头",
             "In contemporary society, [话题] has attracted widespread attention. This essay will explore [分析角度].",
             "灵活适配各类题目"],
        ],
        widths=[3, 8.5, 4.5],
    )
    doc.add_page_break()

    # S6: Body paragraph
    heading(doc, 6, "主体段论证句型", "Task 2 · Body Paragraph")
    body(doc, "主体段是展示逻辑思维和语言能力的核心。以下句型按论证功能分类：")
    make_table(doc,
        ["功能", "句型模板", "适用场景"],
        [
            ["提出论点",
             "One of the primary reasons is that [论点]. / A key argument in favour of this is that [论点].",
             "段落首句"],
            ["举例论证",
             "For instance / For example, [具体例子]. / A case in point is [例子].",
             "支撑论点"],
            ["数据/权威",
             "According to recent studies / research, [论据]. / Statistics suggest that [数据].",
             "增强说服力"],
            ["解释原因",
             "This is primarily because [原因]. / The main reason for this is that [原因].",
             "因果论证"],
            ["解释结果",
             "As a result / Consequently, [结果]. / This leads to / gives rise to [结果].",
             "因果链条"],
            ["让步转折",
             "Admittedly / While it is true that [让步], [转折]. / Despite this, [反驳].",
             "展示批判性思维"],
            ["对比论证",
             "In contrast / By comparison, [对比内容]. / Unlike [A], [B] tends to [差异].",
             "比较分析"],
            ["深入分析",
             "What this means in practice is that [分析]. / In other words, [进一步解释].",
             "展开论述"],
            ["类比推理",
             "This is analogous to [类比]. / Similarly, in the context of [领域], [类比论证].",
             "增强理解"],
            ["假设论证",
             "If [条件], then [结果]. / Without [条件], it is likely that [后果].",
             "假设推理"],
            ["强调重要性",
             "It is worth noting that [要点]. / What is particularly significant is that [要点].",
             "突出关键信息"],
            ["总结段落",
             "Therefore / Thus, it is evident that [段落小结].",
             "段末总结"],
        ],
        widths=[2.5, 8, 5.5],
    )
    doc.add_page_break()

    # S7: Cohesive devices
    heading(doc, 7, "高分衔接与过渡", "Cohesive Devices")
    body(doc, "衔接与连贯（Coherence & Cohesion）占写作评分的 25%。以下是常用的逻辑连接词与短语：")
    make_table(doc,
        ["逻辑关系", "衔接词 / 短语", "例句"],
        [
            ["递进", "Furthermore / Moreover / In addition / Besides",
             "Moreover, technology has made education more accessible."],
            ["对比", "However / In contrast / On the other hand / Conversely",
             "However, not everyone benefits equally from this trend."],
            ["因果", "Therefore / Consequently / As a result / Thus / Hence",
             "Consequently, many young people struggle to find housing."],
            ["举例", "For instance / For example / Such as / A case in point",
             "For instance, countries like Denmark have invested heavily in renewables."],
            ["让步", "Although / Despite / Notwithstanding / Granted",
             "Although costs are high, the long-term benefits are considerable."],
            ["强调", "Indeed / In fact / Notably / Significantly",
             "Indeed, the evidence strongly supports this conclusion."],
            ["条件", "Provided that / As long as / Unless / On condition that",
             "Provided that funding is available, the project will succeed."],
            ["总结", "In conclusion / To sum up / Overall / In summary",
             "In conclusion, the advantages clearly outweigh the disadvantages."],
            ["顺序", "Firstly / Secondly / Finally / Subsequently",
             "Firstly, it is important to consider the economic impact."],
            ["类比", "Similarly / Likewise / In the same way / By analogy",
             "Similarly, other European countries have adopted this approach."],
            ["转换话题", "Turning to / With regard to / As for / In terms of",
             "Turning to the environmental impact, the situation is more complex."],
            ["观点引入", "From my perspective / In my view / It seems to me / I believe",
             "From my perspective, education is the key to solving this issue."],
        ],
        widths=[2.5, 7.5, 6],
    )
    doc.add_page_break()

    # S8: Conclusions
    heading(doc, 8, "结尾段模板", "Task 2 · Conclusion")
    body(doc, "结尾段应总结全文观点，不引入新论点。长度 2-3 句即可：")
    make_table(doc,
        ["题型", "结尾段模板"],
        [
            ["同意/不同意",
             "In conclusion, I firmly [agree/disagree] that [重述立场]. While [让步], [最终观点]."],
            ["讨论双方",
             "In conclusion, while both sides of the argument have merit, I believe that [你的观点]. Ultimately, [总结]."],
            ["优缺点",
             "To sum up, although [优点/缺点], the [advantages/disadvantages] are more significant. Therefore, [建议/展望]."],
            ["问题/方案",
             "In summary, [问题] is a serious issue that requires urgent attention. By [方案], it is possible to [预期效果]."],
            ["积极/消极",
             "In conclusion, while [话题] brings certain challenges, I believe it is largely a [positive/negative] development because [核心原因]."],
            ["报告类",
             "In conclusion, the main factors contributing to [现象] are [原因总结]. Addressing this issue will require [方案总结]."],
            ["通用万能结尾",
             "To conclude, [话题] is a multifaceted issue. The most effective approach would be to [建议], thereby ensuring [积极结果]."],
        ],
        widths=[3.5, 12.5],
    )

    # BACK COVER
    doc.add_page_break()
    for _ in range(5): s=doc.add_paragraph(); _sp(s,0,0)
    _hr(doc,"2F8E87",8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"睿然国际教育",20,True,TEAL); _sp(p,160,80)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p,"用科技赋能教育，让过程被看见，让结果自然发生。",11,False,GRAY_TEXT,i=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,120,0)
    _r(p,"© 睿然国际教育  ·  内部教学资料  ·  版权所有",9,False,GRAY_TEXT)
    _hr(doc,"2F8E87",8)

    doc.save(OUT)
    print(f"✅  Saved → {OUT}")

if __name__=="__main__":
    build()
