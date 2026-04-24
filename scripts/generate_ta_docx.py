#!/usr/bin/env python3
"""Generate a branded Word document for the TA after-class support introduction."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
TEAL = RGBColor(0x2F, 0x8E, 0x87)
TEAL_LIGHT = RGBColor(0xE8, 0xF5, 0xF3)
TEAL_DARK = RGBColor(0x1A, 0x5C, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
OUT_PATH = os.path.join(OUT_DIR, "睿然国际教育_课后教辅体系介绍.docx")


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "2F8E87")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pPr = para._p.get_or_add_pPr()
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}"'
        + (f' w:line="{line}" w:lineRule="auto"' if line else '')
        + '/>'
    )
    # Remove existing spacing
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing)


def add_horizontal_line(doc, color="2F8E87", thickness=6):
    """Add a colored horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    set_paragraph_spacing(p, before=0, after=100)
    return p


def add_section_number(para, number, text):
    """Add a styled section number + title."""
    run_num = para.add_run(f" {number} ")
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = WHITE
    run_num.font.bold = True

    run_text = para.add_run(f"  {text}")
    run_text.font.size = Pt(16)
    run_text.font.color.rgb = TEAL_DARK
    run_text.font.bold = True


def create_branded_table(doc, headers, rows, col_widths=None):
    """Create a branded table with teal header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "2F8E87")

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle bold markers
            if cell_text.startswith("**") and cell_text.endswith("**"):
                run = p.add_run(cell_text[2:-2])
                run.font.bold = True
                run.font.color.rgb = TEAL_DARK
            else:
                run = p.add_run(cell_text)
                run.font.color.rgb = DARK_TEXT
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternating row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F0FAF9")
            set_paragraph_spacing(p, before=60, after=60)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    # Remove default borders and add subtle ones
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)

    return table


def add_body_text(doc, text, bold_segments=None):
    """Add body text with optional bold segments."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)

    # Parse bold markers **text**
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(10.5)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i % 2 == 1:  # odd parts are bold
            run.font.bold = True
            run.font.color.rgb = TEAL_DARK
        else:
            run.font.color.rgb = DARK_TEXT
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)

    # Parse bold markers
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i % 2 == 1:
            run.font.bold = True
            run.font.color.rgb = TEAL_DARK
        else:
            run.font.color.rgb = DARK_TEXT
    return p


def add_callout_box(doc, text):
    """Add a callout/quote box with teal left border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1.5)

    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="2F8E87"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5F3" w:val="clear"/>')
    pPr.append(shading)

    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = TEAL_DARK
    run.font.italic = True
    return p


def add_section_heading(doc, number, title):
    """Add a section heading with number badge."""
    # Add some spacing
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=100, after=0)
    spacer.paragraph_format.space_after = Pt(0)

    # Create a 1-row, 2-col table for the heading badge
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(14.5)

    # Number badge cell
    cell0 = table.rows[0].cells[0]
    cell0.text = ""
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(number)
    run0.font.size = Pt(14)
    run0.font.bold = True
    run0.font.color.rgb = WHITE
    run0.font.name = "Arial"
    cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell0, "2F8E87")

    # Title cell
    cell1 = table.rows[0].cells[1]
    cell1.text = ""
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run(f"  {title}")
    run1.font.size = Pt(16)
    run1.font.bold = True
    run1.font.color.rgb = TEAL_DARK
    run1.font.name = "微软雅黑"
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell1, "E8F5F3")

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)

    # Thin teal bottom line after heading
    add_horizontal_line(doc, "2F8E87", 4)


def add_sub_heading(doc, text):
    """Add a sub-heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── COVER / TITLE SECTION ──
    # Add spacing before title
    for _ in range(4):
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before=0, after=0)

    # Brand name
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_brand.add_run("睿然国际教育")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_paragraph_spacing(p_brand, before=0, after=120)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("RUIRAN INTERNATIONAL EDUCATION")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "Arial"
    run.font.letter_spacing = Pt(3)
    set_paragraph_spacing(p_sub, before=0, after=300)

    # Decorative line
    add_horizontal_line(doc, "2F8E87", 8)

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("课后教辅体系")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = TEAL_DARK
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_paragraph_spacing(p_title, before=200, after=80)

    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title2.add_run("一对一助教陪伴式学习")
    run.font.size = Pt(18)
    run.font.color.rgb = TEAL
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_paragraph_spacing(p_title2, before=0, after=200)

    # Tagline
    add_horizontal_line(doc, "2F8E87", 8)

    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tag.add_run("不止于课堂，更关注课后每一步")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.italic = True
    set_paragraph_spacing(p_tag, before=100, after=60)

    p_tag2 = doc.add_paragraph()
    p_tag2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tag2.add_run("专属助教 · 精细批改 · 实时反馈 · 家长同步")
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.bold = True

    # ── PAGE BREAK ──
    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 01: Why after-class tutoring matters
    # ══════════════════════════════════════════
    add_section_heading(doc, "01", "为什么课后教辅如此重要？")

    add_body_text(doc,
        "在留学备考（托福、雅思、SAT 等）过程中，课堂教学只是学习闭环的一部分。"
        "真正决定提分效果的，往往是**课后的落实与巩固**。"
    )
    add_body_text(doc, "然而，大多数学生在课后面临这样的困境：")

    for item in [
        "作业做了，但不知道对不对",
        "错题看了，但不知道为什么错",
        "计划有了，但没人监督执行",
        "家长想关心，但信息不对称",
    ]:
        add_bullet(doc, item)

    add_callout_box(doc,
        "睿然课后教辅体系正是为了解决这些问题而设计——为每位学生配备专属助教，"
        "通过系统化的批改、反馈和跟踪机制，让课后学习有人管、有人盯、有人帮。"
    )

    # ══════════════════════════════════════════
    # SECTION 02: Dedicated TA
    # ══════════════════════════════════════════
    add_section_heading(doc, "02", "专属助教，一对一绑定")

    add_body_text(doc,
        "我们为每位学生分配**专属助教（Teaching Assistant）**，实行一对一绑定制度："
    )

    create_branded_table(doc,
        headers=["角色", "职责"],
        rows=[
            ["**主讲教师**", "课堂授课、教学规划、阶段性评估"],
            ["**专属助教（Coach）**", "课后任务跟进、每日作业批改、学习督导"],
            ["**批改助教（Reviewer）**", "专项作业精批、正确率评估、错因分析"],
        ],
        col_widths=[4.5, 11.5],
    )

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=80, after=0)

    add_body_text(doc, "每位助教在系统中与学生建立绑定关系，确保：")
    for item in [
        '**专人负责**：您的孩子不会被"踢皮球"，永远有明确的课后责任人',
        "**持续跟踪**：助教长期跟进同一学生，深入了解学习习惯和薄弱点",
        "**主辅配合**：主讲教师与助教信息互通，教学与辅导无缝衔接",
    ]:
        add_bullet(doc, "✅ " + item)

    # ══════════════════════════════════════════
    # SECTION 03: Daily tasks
    # ══════════════════════════════════════════
    add_section_heading(doc, "03", "每日任务与作业管理")

    add_sub_heading(doc, "精细化的每日学习计划")
    add_body_text(doc,
        "助教每天根据主讲教师的教学进度，为学生制定**个性化的课后任务清单**："
    )
    for item in [
        "**按科目分类**：阅读、听力、写作、口语等模块独立安排",
        "**明确时间规划**：每项任务标注建议用时，培养时间管理能力",
        "**灵活证据要求**：根据任务类型，要求提交文字、图片或录音作为完成凭证",
    ]:
        add_bullet(doc, item)

    add_sub_heading(doc, "学生端实时打卡")
    add_body_text(doc, "学生通过微信小程序完成每日任务：")
    for i, item in enumerate([
        "查看当日任务清单",
        "逐项完成并提交证据（拍照上传作业、录音提交口语等）",
        "记录实际用时，系统自动统计学习时长",
    ], 1):
        add_bullet(doc, f"{i}. {item}")

    # ══════════════════════════════════════════
    # SECTION 04: Grading & Feedback (CORE)
    # ══════════════════════════════════════════
    add_section_heading(doc, "04", "作业批改与反馈机制（核心亮点）")

    add_callout_box(doc,
        "这是睿然课后教辅体系的核心竞争力——每一份作业都会经过助教的认真批改和反馈。"
    )

    add_sub_heading(doc, "三级评审状态")
    create_branded_table(doc,
        headers=["状态", "含义"],
        rows=[
            ["**✅ 通过（Approved）**", "作业完成质量达标，可进入下一阶段"],
            ["**⚠️ 部分通过（Partial）**", "基本完成，但存在需改进的地方，附带具体建议"],
            ["**❌ 未通过（Rejected）**", "需要重新完成或订正，助教会说明原因"],
        ],
        col_widths=[5.5, 10.5],
    )

    add_sub_heading(doc, "多维度反馈内容")
    add_body_text(doc, "助教的反馈不是简单的对错判断，而是**有深度的个性化指导**：")
    for item in [
        "**正确率评估**：量化作业准确率，追踪进步趋势",
        "**文字批注**：针对具体错误给出解析和改进建议",
        "**图片批改**：在学生作业图片上直接圈点勾画，标注问题所在",
        "**语音点评**：录制语音反馈，如同老师面对面辅导",
    ]:
        add_bullet(doc, "📌 " + item)

    add_sub_heading(doc, "完整的审批追踪")
    add_body_text(doc, "系统记录每一次批改的完整轨迹：")
    for item in [
        "批改人、批改时间、状态变更历史",
        "助教评语永久留存，可随时回看",
        "家长可在小程序中直接查看所有反馈记录",
    ]:
        add_bullet(doc, item)

    # ══════════════════════════════════════════
    # SECTION 05: Post-class feedback report
    # ══════════════════════════════════════════
    add_section_heading(doc, "05", "课后反馈报告（每节课后自动推送）")

    add_body_text(doc,
        "每节课结束后，授课教师通过系统提交**结构化课后反馈**，内容涵盖三大板块："
    )

    create_branded_table(doc,
        headers=["板块", "内容"],
        rows=[
            ["**作业完成情况**", "课前作业的完成度、准确率、存在的问题"],
            ["**课堂表现及问题**", "课堂专注度、互动表现、关键知识点掌握情况"],
            ["**建议及作业**", "课后建议与下次课前作业安排"],
        ],
        col_widths=[4.5, 11.5],
    )

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=60, after=0)

    add_body_text(doc, "反馈提交后：")
    for item in [
        "📲 **微信即时推送**：家长第一时间收到通知",
        "📱 **小程序随时查看**：所有历史反馈永久保存，支持回看",
        "🖼 **支持图片附件**：教师可附上课堂板书、标注等图片",
    ]:
        add_bullet(doc, item)

    add_callout_box(doc,
        '家长不用再追着老师问"今天上课怎么样"——打开手机，反馈已经在那里了。'
    )

    # ══════════════════════════════════════════
    # SECTION 06: Data-driven tracking
    # ══════════════════════════════════════════
    add_section_heading(doc, "06", "数据驱动的学习追踪")

    add_body_text(doc,
        '课后教辅不只是"布置-批改"的简单循环。我们的系统提供全方位的数据追踪：'
    )
    for item in [
        "📈 **每日完成率统计**：一目了然地掌握孩子每天的任务执行情况",
        "📊 **7日学习趋势图**：用可视化图表呈现一周的学习状态波动",
        "🍰 **学科时间分布**：清晰展示各科目投入比例，避免偏科",
        "⏱ **学习时长记录**：精确到每项任务的实际用时",
        "📋 **历史成绩对比**：阶段性考试成绩录入系统，关联学习投入看产出",
    ]:
        add_bullet(doc, item)

    # ══════════════════════════════════════════
    # SECTION 07: Parent participation
    # ══════════════════════════════════════════
    add_section_heading(doc, "07", "家长如何参与？")

    add_body_text(doc, "我们相信，最好的教育是**家校协同**。通过睿然学习追踪小程序，家长可以：")

    create_branded_table(doc,
        headers=["功能", "说明"],
        rows=[
            ["**实时查看任务进度**", "孩子今天要做什么、做了多少、还剩什么"],
            ["**接收课后反馈推送**", "每节课后自动推送教师反馈到微信"],
            ["**查看作业批改详情**", "助教的文字批注、图片标注、语音点评"],
            ["**浏览学习数据看板**", "完成率趋势、学科分布、时间投入"],
            ["**多孩子一键切换**", "多个孩子独立档案，一个微信号统一管理"],
        ],
        col_widths=[5, 11],
    )

    # ══════════════════════════════════════════
    # SECTION 08: Commitments
    # ══════════════════════════════════════════
    add_section_heading(doc, "08", "我们的承诺")

    create_branded_table(doc,
        headers=["承诺", "详情"],
        rows=[
            ["**专人负责制**", "每位学生有明确的助教负责人，杜绝推诿"],
            ["**当日反馈制**", "作业提交后，助教在当日内完成批改与反馈"],
            ["**全程可追溯**", "所有批改记录、反馈内容永久保存，透明可查"],
            ["**家长知情权**", "课后反馈实时推送，家长随时掌握学习动态"],
            ["**数据安全**", "学生数据严格保密，仅授权人员可访问"],
        ],
        col_widths=[4.5, 11.5],
    )

    # ══════════════════════════════════════════
    # SECTION 09: Getting started
    # ══════════════════════════════════════════
    add_section_heading(doc, "09", "如何开始？")

    add_body_text(doc, "只需三步，即可为您的孩子开启专属课后教辅服务：")

    for i, (title, desc) in enumerate([
        ("咨询报名", "联系睿然课程顾问，了解适合的课程方案"),
        ("匹配助教", "根据学生的科目、水平和需求，分配专属助教"),
        ("开启追踪", "家长微信扫码进入小程序，实时掌握学习全貌"),
    ], 1):
        add_bullet(doc, f"**第{i}步 · {title}**：{desc}")

    # ── FOOTER / CLOSING ──
    doc.add_paragraph()
    add_horizontal_line(doc, "2F8E87", 8)

    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_close.paragraph_format.space_before = Pt(20)
    run = p_close.add_run("睿然国际教育")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p_slogan = doc.add_paragraph()
    p_slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_slogan.add_run('用科技赋能教育，让课后不再是"无人区"')
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.italic = True

    p_slogan2 = doc.add_paragraph()
    p_slogan2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_slogan2.add_run("每一份作业有回应，每一个问题有解答，每一天的努力被看见。")
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.italic = True

    # ── Save ──
    doc.save(OUT_PATH)
    print(f"✅ Document saved to: {OUT_PATH}")


if __name__ == "__main__":
    build_document()
