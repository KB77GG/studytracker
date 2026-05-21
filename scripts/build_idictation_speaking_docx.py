#!/usr/bin/env python3
"""Build a Word document from an idictation speaking export JSON."""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def clean_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\ufeff", "").replace("\u200b", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_number(text: str) -> str:
    return re.sub(r"^\s*\d+\s*[.)、]\s*", "", clean_text(text)).strip()


def load_payload(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if "part1" in payload and "part23" in payload:
        return payload
    if "questions" not in payload:
        raise ValueError("JSON must contain either part1/part23 or questions.")
    return organize_from_flat(payload)


def organize_from_flat(payload: dict[str, Any]) -> dict[str, Any]:
    part1: dict[str, dict[str, Any]] = {}
    part23: dict[str, dict[str, Any]] = {}
    for row in payload.get("questions") or []:
        material_id = clean_text(row.get("material_id"))
        topic = clean_text(row.get("topic") or row.get("material_title"))
        part = clean_text(row.get("part"))
        issue_id = clean_text(row.get("issue_id"))
        source_date = clean_text(row.get("source_date") or row.get("suggested_time") or row.get("updated_at"))
        question = clean_text(row.get("question"))
        if not material_id or not question:
            continue
        if part == "Part 1":
            bucket = part1.setdefault(
                material_id,
                {"material_id": material_id, "topic": topic, "source_date": source_date, "questions": []},
            )
            bucket["questions"].append({"issue_id": issue_id, "question": strip_number(question), "source_date": source_date})
        else:
            bucket = part23.setdefault(
                material_id,
                {"material_id": material_id, "topic": topic, "source_date": source_date, "part2": None, "part3": []},
            )
            if part == "Part 2":
                bucket["part2"] = {"issue_id": issue_id, "card": question, "source_date": source_date}
            elif part == "Part 3":
                bucket["part3"].append({"issue_id": issue_id, "question": strip_number(question), "source_date": source_date})
    payload["part1"] = list(part1.values())
    payload["part23"] = list(part23.values())
    return payload


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def set_table_columns(table, widths: list[Any]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width.inches * 1440)))
        tbl_grid.append(grid_col)


def set_font(run, name: str = "Aptos", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_multiline_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    lines = clean_text(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_font(run)
    return paragraph


def add_meta_paragraph(document: Document, payload: dict[str, Any], source_path: Path) -> None:
    counts = payload.get("counts") or {}
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")
    items = [
        f"Generated: {generated}",
        f"Source file: {source_path.name}",
        f"Latest source date: {payload.get('latest_source_date') or '-'}",
        (
            "Counts: "
            f"Part 1 {counts.get('part1_topics', len(payload.get('part1') or []))} topics / "
            f"{counts.get('part1_questions', sum(len(t.get('questions') or []) for t in payload.get('part1') or []))} questions; "
            f"Part 2+3 {counts.get('part23_topics', len(payload.get('part23') or []))} topics."
        ),
    ]
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(" | ".join(items))
    set_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 20, RGBColor(25, 44, 70)),
        ("Heading 1", 15, RGBColor(25, 44, 70)),
        ("Heading 2", 12, RGBColor(35, 65, 100)),
        ("Heading 3", 10.5, RGBColor(35, 65, 100)),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_part1(document: Document, topics: list[dict[str, Any]]) -> None:
    document.add_heading("Part 1", level=1)
    for topic_index, topic in enumerate(topics, 1):
        source_date = clean_text(topic.get("source_date"))
        title = f"{topic_index}. {clean_text(topic.get('topic'))}"
        if source_date:
            title += f" ({source_date})"
        document.add_heading(title, level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        set_table_columns(table, [Inches(0.35), Inches(6.45)])
        header = table.rows[0].cells
        set_cell_text(header[0], "#", bold=True)
        set_cell_text(header[1], "Question", bold=True)
        set_cell_shading(header[0], "D9EAF7")
        set_cell_shading(header[1], "D9EAF7")
        for question_index, question in enumerate(topic.get("questions") or [], 1):
            cells = table.add_row().cells
            set_cell_text(cells[0], str(question_index))
            set_cell_text(cells[1], clean_text(question.get("question")))
        document.add_paragraph()


def add_part23(document: Document, topics: list[dict[str, Any]]) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Part 2 and Part 3", level=1)
    for topic_index, topic in enumerate(topics, 1):
        source_date = clean_text(topic.get("source_date"))
        title = f"{topic_index}. {clean_text(topic.get('topic'))}"
        if source_date:
            title += f" ({source_date})"
        document.add_heading(title, level=2)

        part2 = topic.get("part2") or {}
        if part2:
            document.add_heading("Part 2 Card", level=3)
            paragraph = add_multiline_paragraph(document, clean_text(part2.get("card")))
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.space_after = Pt(6)

        part3 = topic.get("part3") or []
        if part3:
            document.add_heading("Part 3 Questions", level=3)
            for question_index, question in enumerate(part3, 1):
                paragraph = document.add_paragraph(style="List Number")
                run = paragraph.add_run(clean_text(question.get("question")))
                set_font(run)
                run.font.size = Pt(10.5)
        document.add_paragraph()


def build_docx(source_path: Path, output_path: Path) -> None:
    payload = load_payload(source_path)
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IELTS Speaking Question Bank")
    set_font(run)
    add_meta_paragraph(document, payload, source_path)

    add_part1(document, payload.get("part1") or [])
    add_part23(document, payload.get("part23") or [])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a DOCX from idictation speaking export JSON.")
    parser.add_argument("--input", default="data/idictation_speaking/organized/ielts_speaking_questions.json")
    parser.add_argument("--output", default="output/doc/ielts_speaking_question_bank.docx")
    args = parser.parse_args()

    build_docx(Path(args.input), Path(args.output))
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
