#!/usr/bin/env python3
"""Build the 2026.5-8 IELTS speaking question bank DOCX from a live export."""

from __future__ import annotations

import argparse
import importlib.util
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BRAND = RGBColor(47, 142, 135)
BRAND_DARK = RGBColor(39, 124, 120)
BRAND_LIGHT = RGBColor(64, 166, 160)
BRAND_SOFT = "E5F4F2"
BRAND_TINT = "F2FBF9"
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
BRAND_MARK = ROOT / "static" / "brand" / "sagepath-mark.png"

PEOPLE_IDS = {
    "1006", "1010", "1021", "1023", "1024", "1036",
    "406", "546", "634", "909", "993", "995",
}
PLACE_IDS = {"1007", "1017", "1019", "1020", "1032"}
THING_IDS = {"1018", "1026", "1028", "324", "370", "426", "760", "988", "997", "998"}


def clean_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    text = re.sub(r"(?m)^ls\b", "Is", text)
    text = text.replace("stillremember", "still remember")
    text = text.replace("f or your family", "for your family")
    return text


def topic_map(topics: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {clean_text(topic.get("topic")): topic for topic in topics}


def dedupe_questions(questions: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for question in questions:
        question = clean_text(question)
        key = question.lower()
        if not question or key in seen:
            continue
        seen.add(key)
        unique.append(question)
    return unique


def question_texts(topic: dict[str, Any] | None) -> list[str]:
    if not topic:
        return []
    return [clean_text(item.get("question")) for item in topic.get("questions") or []]


def load_export_helpers():
    module_path = ROOT / "scripts" / "export_idictation_speaking.py"
    spec = importlib.util.spec_from_file_location("idictation_export", module_path)
    if not spec or not spec.loader:
        raise RuntimeError(f"Cannot load {module_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def material_metadata(raw_path: Path) -> dict[str, dict[str, Any]]:
    if not raw_path.exists():
        return {}
    helper = load_export_helpers()
    raw = json.loads(raw_path.read_text(encoding="utf-8"))
    metadata: dict[str, dict[str, Any]] = {}
    for source_part in ("part1", "part23"):
        for entry in raw.get("lists", {}).get(source_part, []):
            for item in helper.find_items(helper.unwrap_values(entry.get("response") or {})):
                material = helper.normalize_material(item, source_part)
                mid = clean_text(material.get("material_id"))
                if mid and mid not in metadata:
                    metadata[mid] = {
                        "order": len(metadata),
                        "is_new": clean_text(item.get("mkt_new_topic")) == "238",
                    }
    return metadata


def material_order(metadata: dict[str, dict[str, Any]]) -> dict[str, int]:
    return {mid: int(item.get("order", 10_000)) for mid, item in metadata.items()}


def is_new_topic(topic: dict[str, Any], metadata: dict[str, dict[str, Any]]) -> bool:
    mid = clean_text(topic.get("material_id"))
    return bool((metadata.get(mid) or {}).get("is_new"))


def heading_title(index: int, title: str, is_new: bool) -> str:
    suffix = " 【新题】" if is_new else ""
    return f"{index}. {title}{suffix}"


def sort_topics_new_first(
    topics: list[dict[str, Any]],
    metadata: dict[str, dict[str, Any]],
    order: dict[str, int],
) -> list[dict[str, Any]]:
    return sorted(
        topics,
        key=lambda topic: (
            0 if is_new_topic(topic, metadata) else 1,
            order.get(clean_text(topic.get("material_id")), 10_000),
        ),
    )


def build_part1_topics(
    payload: dict[str, Any],
    order: dict[str, int],
    metadata: dict[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    topics = topic_map(payload.get("part1") or [])
    work_studies = question_texts(topics.get("Work/studies"))
    general = work_studies[:1]
    work_questions = general + work_studies[1:10]
    study_questions = general + work_studies[10:]

    home_questions = (
        question_texts(topics.get("Home/Accommodation"))
        + question_texts(topics.get("The area you live in"))
        + question_texts(topics.get("The city you live in"))
    )

    required = [
        {
            "topic": "家乡 / Hometown",
            "source_topic": "Hometown",
            "questions": dedupe_questions(question_texts(topics.get("Hometown"))),
        },
        {
            "topic": "学习 / Study",
            "source_topic": "Work/studies",
            "questions": dedupe_questions(study_questions),
        },
        {
            "topic": "工作 / Work",
            "source_topic": "Work/studies",
            "questions": dedupe_questions(work_questions),
        },
        {
            "topic": "居住地 / Home, Accommodation & Local Area",
            "source_topic": "Home/Accommodation; The area you live in; The city you live in",
            "questions": dedupe_questions(home_questions),
        },
    ]

    consumed = {
        "Work/studies",
        "Home/Accommodation",
        "Hometown",
        "The area you live in",
        "The city you live in",
    }
    rest = [
        topic
        for topic in payload.get("part1") or []
        if clean_text(topic.get("topic")) not in consumed and question_texts(topic)
    ]
    rest = sort_topics_new_first(rest, metadata, order)
    return required, rest


def real_part23_topics(payload: dict[str, Any], order: dict[str, int]) -> list[dict[str, Any]]:
    topics = []
    for topic in payload.get("part23") or []:
        part2 = topic.get("part2") or {}
        card = clean_text(part2.get("card"))
        if not card.lower().startswith("describe "):
            continue
        topics.append(topic)
    topics.sort(key=lambda topic: order.get(clean_text(topic.get("material_id")), 10_000))
    return topics


def categorize(topic: dict[str, Any]) -> str:
    mid = clean_text(topic.get("material_id"))
    if mid in PEOPLE_IDS:
        return "人物类 People"
    if mid in PLACE_IDS:
        return "地点类 Places"
    if mid in THING_IDS:
        return "物品类 Objects / Things"
    return "事件类 Events"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, RGBColor(31, 41, 55))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    set_style_font(styles["Title"], 24, INK, True)
    styles["Title"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 1"], 16, BRAND, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], 13, BRAND_DARK, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    set_style_font(styles["Heading 3"], 12, RGBColor(22, 78, 74), True)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("睿然国际教育 · IELTS Speaking 2026.5-8  |  Page ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(footer, "PAGE")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_bottom_border(paragraph, color: str, size: int = 12, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def table_pr(table, widths_dxa: list[int], fill_header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, width in CELL_MARGINS_DXA.items():
        el = cell_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            cell_mar.append(el)
        el.set(qn("w:w"), str(width))
        el.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
    if fill_header and table.rows:
        for cell in table.rows[0].cells:
            shade_cell(cell, BRAND_SOFT)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 9.5, align: str = "left") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    for index, line in enumerate(clean_text(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size, bold)


def add_small_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, 9, color=RGBColor(85, 85, 85))


def add_questions_table(document: Document, questions: list[str]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table_pr(table, [520, CONTENT_WIDTH_DXA - 520])
    set_cell(table.rows[0].cells[0], "#", bold=True, align="center")
    set_cell(table.rows[0].cells[1], "Question", bold=True)
    for index, question in enumerate(questions, 1):
        row = table.add_row()
        set_cell(row.cells[0], str(index), align="center")
        set_cell(row.cells[1], question)
    document.add_paragraph()


def add_card_table(document: Document, card: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_pr(table, [CONTENT_WIDTH_DXA], fill_header=False)
    shade_cell(table.rows[0].cells[0], BRAND_TINT)
    set_cell(table.rows[0].cells[0], card, size=10)


def add_title(
    document: Document,
    payload: dict[str, Any],
    source_url: str,
    required: list[dict[str, Any]],
    rest: list[dict[str, Any]],
    part23: list[dict[str, Any]],
) -> None:
    if BRAND_MARK.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.paragraph_format.space_after = Pt(2)
        logo.add_run().add_picture(str(BRAND_MARK), width=Inches(0.58))

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_after = Pt(2)
    run = brand.add_run("睿然国际教育")
    set_run_font(run, 12, True, BRAND_DARK)
    run = brand.add_run("  |  Sage Path IELTS Speaking Studio")
    set_run_font(run, 9.5, False, MUTED)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IELTS Speaking Question Bank")
    set_run_font(run, 25, True, INK)
    add_bottom_border(title, "2F8E87", size=14, space=10)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("2026.5-8月新版题库 · 新题优先排序 · Part 1 / Part 2 / Part 3")
    set_run_font(run, 13, True, BRAND_DARK)

    part1_question_count = sum(len(topic.get("questions") or []) for topic in required)
    part1_question_count += sum(len(question_texts(topic)) for topic in rest)
    part2_count = len(part23)
    part3_count = sum(len(topic.get("part3") or []) for topic in part23)
    add_small_note(
        document,
        "题库概览："
        f"Part 1 {len(required) + len(rest)} 个话题 / {part1_question_count} 题；"
        f"Part 2 题卡 {part2_count} 张；Part 3 延伸题 {part3_count} 题。",
    )


def add_part1(
    document: Document,
    required: list[dict[str, Any]],
    rest: list[dict[str, Any]],
    metadata: dict[str, dict[str, Any]],
) -> None:
    document.add_heading("Part 1", level=1)
    document.add_heading("必考题 Required Topics", level=2)
    for index, topic in enumerate(required, 1):
        heading = document.add_heading(f"{index}. {topic['topic']}", level=3)
        heading.paragraph_format.keep_with_next = True
        add_questions_table(document, topic["questions"])

    document.add_heading("当季题 Other Current Topics", level=2)
    for index, topic in enumerate(rest, 1):
        title = clean_text(topic.get("topic"))
        heading = document.add_heading(heading_title(index, title, is_new_topic(topic, metadata)), level=3)
        heading.paragraph_format.keep_with_next = True
        add_questions_table(document, question_texts(topic))


def add_part23(document: Document, topics: list[dict[str, Any]], metadata: dict[str, dict[str, Any]]) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Part 2 & Part 3", level=1)

    categories = ["人物类 People", "地点类 Places", "物品类 Objects / Things", "事件类 Events"]
    grouped = {category: [] for category in categories}
    for topic in topics:
        grouped[categorize(topic)].append(topic)

    for category in categories:
        document.add_heading(category, level=2)
        grouped[category] = sort_topics_new_first(grouped[category], metadata, {
            mid: int(item.get("order", 10_000)) for mid, item in metadata.items()
        })
        for index, topic in enumerate(grouped[category], 1):
            title = clean_text(topic.get("topic"))
            heading = document.add_heading(heading_title(index, title, is_new_topic(topic, metadata)), level=3)
            heading.paragraph_format.keep_with_next = True

            part2 = topic.get("part2") or {}
            label = document.add_paragraph()
            label.paragraph_format.keep_with_next = True
            label.paragraph_format.space_after = Pt(4)
            run = label.add_run("Part 2 Cue Card")
            set_run_font(run, 10.5, True, RGBColor(31, 77, 120))
            add_card_table(document, clean_text(part2.get("card")))

            part3 = [clean_text(item.get("question")) for item in topic.get("part3") or []]
            if part3:
                label = document.add_paragraph()
                label.paragraph_format.keep_with_next = True
                label.paragraph_format.space_before = Pt(6)
                label.paragraph_format.space_after = Pt(4)
                run = label.add_run("Part 3 Questions")
                set_run_font(run, 10.5, True, RGBColor(31, 77, 120))
                add_questions_table(document, part3)
            else:
                document.add_paragraph()


def build_docx(input_path: Path, raw_path: Path, output_path: Path, source_url: str) -> None:
    payload = json.loads(input_path.read_text(encoding="utf-8"))
    metadata = material_metadata(raw_path)
    order = material_order(metadata)
    required, rest = build_part1_topics(payload, order, metadata)
    part23 = real_part23_topics(payload, order)

    doc = Document()
    configure_document(doc)
    add_title(doc, payload, source_url, required, rest, part23)
    add_part1(doc, required, rest, metadata)
    add_part23(doc, part23, metadata)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the seasonal IELTS speaking question bank DOCX.")
    parser.add_argument("--input", default="data/idictation_speaking_live_20260515/organized/ielts_speaking_questions.json")
    parser.add_argument("--raw", default="data/idictation_speaking_live_20260515/speaking_materials.raw.json")
    parser.add_argument("--output", default="/Users/zhouxin/Desktop/IELTS_Speaking_QuestionBank_2026.5-8新版.docx")
    parser.add_argument("--source-url", default="https://www.idictation.cn/main/book")
    args = parser.parse_args()
    build_docx(Path(args.input), Path(args.raw), Path(args.output), args.source_url)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
