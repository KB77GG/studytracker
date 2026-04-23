#!/usr/bin/env python3
"""Generate branded TOEFL Core Vocabulary document – same style as IELTS vocab booklet."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

TEAL=RGBColor(0x2F,0x8E,0x87); TEAL_DARK=RGBColor(0x1A,0x5C,0x57)
WHITE=RGBColor(0xFF,0xFF,0xFF); DARK=RGBColor(0x2D,0x2D,0x2D)
GRAY=RGBColor(0x66,0x66,0x66); ORANGE=RGBColor(0xE8,0x7C,0x2A)
B80=RGBColor(0xE6,0x8A,0x00); B90=TEAL; B100=RGBColor(0x7B,0x1F,0xA2)
BG80="FFF8E1"; BG90="E8F5F3"; BG100="EDE7F6"

OUT=os.path.join(os.path.dirname(os.path.dirname(__file__)),"docs",
    "睿然国际教育_托福核心词汇手册.docx")

# ── helpers ──
def _sh(c,h): c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}" w:val="clear"/>'))
def _sp(p,b=0,a=0):
    pPr=p._p.get_or_add_pPr()
    for o in pPr.findall(qn('w:spacing')): pPr.remove(o)
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="{b}" w:after="{a}"/>'))
def _hr(d,c="2F8E87",s=6):
    p=d.add_paragraph(); p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{s}" w:space="1" w:color="{c}"/></w:pBdr>')); _sp(p,0,100)
def _r(p,t,sz=10.5,b=False,c=DARK,i=False,f="微软雅黑"):
    r=p.add_run(t);r.font.size=Pt(sz);r.font.bold=b;r.font.color.rgb=c;r.font.name=f
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    if i: r.font.italic=True
    return r
def _nb(t):
    tP=t._tbl.tblPr
    if tP is None: tP=parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    b=parse_xml(f'<w:tblBorders {nsdecls("w")}>'+''.join(f'<w:{e} w:val="none" w:sz="0" w:space="0" w:color="auto"/>' for e in("top","left","bottom","right","insideH","insideV"))+'</w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(b)
def _tb(t):
    tP=t._tbl.tblPr
    if tP is None: tP=parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    b=parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:left w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:right w:val="single" w:sz="4" w:space="0" w:color="2F8E87"/><w:insideH w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/><w:insideV w:val="single" w:sz="2" w:space="0" w:color="BDDBD8"/></w:tblBorders>')
    for o in tP.findall(qn('w:tblBorders')): tP.remove(o)
    tP.append(b)

def hd(d,n,cn,en=""):
    s=d.add_paragraph();_sp(s,140,0)
    t=d.add_table(rows=1,cols=2);t.autofit=False;t.columns[0].width=Cm(1.5);t.columns[1].width=Cm(24)
    c0=t.rows[0].cells[0];c0.text="";p0=c0.paragraphs[0];p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _r(p0,str(n).zfill(2),13,True,WHITE,f="Arial");c0.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c0,"2F8E87")
    c1=t.rows[0].cells[1];c1.text="";p1=c1.paragraphs[0]
    _r(p1,f"  {cn}",15,True,TEAL_DARK)
    if en: _r(p1,f"  {en}",11,False,GRAY,f="Arial")
    c1.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c1,"E8F5F3");_nb(t);_hr(d,"2F8E87",4)

BAND_BG={"80+":BG80,"90+":BG90,"100+":BG100}
BAND_C={"80+":B80,"90+":B90,"100+":B100}

def vtable(d,entries):
    hdr=["#","Word / Phrase","词性","中文释义","TPO 例句","板块"]
    cw=[0.7,3.2,0.8,2.3,6.5,2.5]
    bo={"80+":0,"90+":1,"100+":2}
    es=sorted(entries,key=lambda e:(bo.get(e[5],9),e[0].lower()))
    t=d.add_table(rows=1+len(es),cols=6);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    for i,h in enumerate(hdr):
        c=t.rows[0].cells[i];c.text="";p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _r(p,h,9,True,WHITE);c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c,"2F8E87");c.width=Cm(cw[i])
    for idx,(w,pos,cn,ex,sec,band) in enumerate(es):
        row=t.rows[idx+1]; bg=BAND_BG.get(band,"FFFFFF"); bc=BAND_C.get(band,DARK)
        vals=[str(idx+1),w,pos,cn,ex,sec]
        als=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
             WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER]
        for ci,v in enumerate(vals):
            c=row.cells[ci];c.text="";c.width=Cm(cw[ci]);p=c.paragraphs[0];p.alignment=als[ci];_sp(p,30,30)
            if ci==1: _r(p,v,9.5,True,TEAL_DARK,f="Arial")
            elif ci==4: _r(p,v,8.5,False,GRAY,i=True,f="Arial")
            elif ci==5: _r(p,v,8.5,True,bc,f="Arial")
            else: _r(p,v,9,False,DARK)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER;_sh(c,bg)
    _tb(t)
    pl=d.add_paragraph();pl.paragraph_format.space_before=Pt(6)
    _r(pl,"  ■",9,True,B80,f="Arial");_r(pl," 80+ 基础   ",8,False,GRAY)
    _r(pl,"■",9,True,B90,f="Arial");_r(pl," 90+ 进阶   ",8,False,GRAY)
    _r(pl,"■",9,True,B100,f="Arial");_r(pl," 100+ 高阶",8,False,GRAY)

# ── VOCAB DATA ──
V={
"生物与生态 Biology & Ecology":[
    ("organism","n.","有机体","All living organisms require energy to maintain their biological functions.","R L","80+"),
    ("habitat","n.","栖息地","The destruction of natural habitats threatens countless species.","R L","80+"),
    ("predator","n.","捕食者","Predators play a crucial role in maintaining ecological balance.","R L","80+"),
    ("species","n.","物种","Over 8 million species inhabit the Earth, many still undiscovered.","R L","80+"),
    ("migration","n.","迁徙","Seasonal migration allows birds to exploit different food sources.","R L","80+"),
    ("ecosystem","n.","生态系统","Coral reefs are among the most diverse ecosystems on the planet.","R L","90+"),
    ("photosynthesis","n.","光合作用","Photosynthesis converts sunlight into chemical energy stored in glucose.","R L","90+"),
    ("adaptation","n.","适应","Desert plants show remarkable adaptations to conserve water.","R L","90+"),
    ("symbiosis","n.","共生","Symbiosis between fungi and plant roots benefits both organisms.","R","90+"),
    ("extinction","n.","灭绝","Mass extinction events have reshaped life on Earth multiple times.","R L","90+"),
    ("vertebrate","n.","脊椎动物","Vertebrates account for only a small fraction of animal species.","R","90+"),
    ("metabolism","n.","新陈代谢","Cold-blooded animals have lower metabolic rates than warm-blooded ones.","R","100+"),
    ("marine","adj.","海洋的","Marine biologists study organisms in oceanic environments.","R L","80+"),
    ("parasite","n.","寄生虫","Parasites often evolve complex strategies to exploit their hosts.","R","100+"),
    ("invertebrate","n.","无脊椎动物","Invertebrates make up over 95% of all known animal species.","R","100+"),
],
"地质与地球科学 Geology & Earth Science":[
    ("erosion","n.","侵蚀","Wind and water erosion shaped the canyon over millions of years.","R L","80+"),
    ("fossil","n.","化石","Fossils provide direct evidence of ancient life forms.","R L","80+"),
    ("volcanic","adj.","火山的","Volcanic eruptions can significantly alter local landscapes.","R L","80+"),
    ("earthquake","n.","地震","Earthquakes occur along fault lines where tectonic plates meet.","R L","80+"),
    ("mineral","n.","矿物","Minerals are naturally occurring inorganic substances with defined compositions.","R L","80+"),
    ("sediment","n.","沉积物","Layers of sediment accumulate on the ocean floor over time.","R L","90+"),
    ("glacier","n.","冰川","Glaciers carved deep valleys during the last Ice Age.","R L","90+"),
    ("tectonic","adj.","构造的","Tectonic plate movement drives continental drift and mountain building.","R","90+"),
    ("crust","n.","地壳","The Earth's crust is thinnest beneath the oceans.","R L","90+"),
    ("precipitation","n.","降水","Precipitation patterns are shifting due to climate change.","R L","90+"),
    ("geothermal","adj.","地热的","Geothermal energy harnesses heat from deep within the Earth.","R","100+"),
    ("stratum","n.","地层","Each stratum of rock tells a story about past environments.","R","100+"),
    ("aquifer","n.","含水层","Underground aquifers supply drinking water to many communities.","R","100+"),
    ("meteorology","n.","气象学","Meteorology uses atmospheric data to forecast weather patterns.","R","100+"),
    ("continental drift","n.","大陆漂移","Continental drift theory explains the movement of landmasses over time.","R","90+"),
],
"天文与宇宙 Astronomy & Space":[
    ("orbit","n./v.","轨道；环绕","The Earth orbits the Sun once every 365.25 days.","R L","80+"),
    ("solar system","n.","太阳系","Our solar system contains eight officially recognised planets.","R L","80+"),
    ("gravity","n.","重力","Gravity governs the motion of planets and stars.","R L","80+"),
    ("satellite","n.","卫星","Artificial satellites are used for communication and observation.","R L","80+"),
    ("telescope","n.","望远镜","The Hubble telescope has captured images of distant galaxies.","R L","80+"),
    ("galaxy","n.","星系","The Milky Way galaxy contains hundreds of billions of stars.","R L","90+"),
    ("asteroid","n.","小行星","An asteroid impact is believed to have caused dinosaur extinction.","R L","90+"),
    ("constellation","n.","星座","Ancient cultures used constellations for navigation and storytelling.","R L","90+"),
    ("nebula","n.","星云","A nebula is a cloud of gas and dust where stars are born.","R","100+"),
    ("supernova","n.","超新星","A supernova explosion can briefly outshine an entire galaxy.","R","100+"),
    ("celestial","adj.","天体的","Celestial navigation relies on the positions of stars and planets.","R","90+"),
    ("cosmos","n.","宇宙","Scientists continue to explore the vast mysteries of the cosmos.","R","90+"),
    ("light-year","n.","光年","The nearest star is about four light-years from Earth.","R L","90+"),
    ("astronomical","adj.","天文学的","Astronomical observations have revealed thousands of exoplanets.","R","100+"),
    ("dwarf planet","n.","矮行星","Pluto was reclassified as a dwarf planet in 2006.","R L","90+"),
],
"历史与考古 History & Archaeology":[
    ("ancient","adj.","古代的","Ancient civilizations developed sophisticated writing systems.","R L","80+"),
    ("civilization","n.","文明","The rise of civilization was linked to agricultural development.","R L","80+"),
    ("monument","n.","纪念碑","The pyramids are enduring monuments to Egyptian engineering.","R L","80+"),
    ("dynasty","n.","朝代","The Ming Dynasty is renowned for its cultural achievements.","R L","80+"),
    ("revolution","n.","革命","The Industrial Revolution transformed manufacturing and society.","R L","80+"),
    ("artifact","n.","文物","Artifacts recovered from the site date back to 3000 BCE.","R L","90+"),
    ("excavation","n.","挖掘","Excavation of the ancient city revealed elaborate drainage systems.","R","90+"),
    ("colonization","n.","殖民","European colonization profoundly altered indigenous societies.","R L","90+"),
    ("indigenous","adj.","本土的","Indigenous peoples had complex social structures before contact.","R L","90+"),
    ("prehistoric","adj.","史前的","Prehistoric cave paintings offer insight into early human culture.","R","90+"),
    ("empire","n.","帝国","The Roman Empire extended across three continents at its peak.","R L","80+"),
    ("archaeology","n.","考古学","Archaeology combines fieldwork with laboratory analysis.","R","90+"),
    ("chronicle","n.","编年史","Medieval chronicles recorded the events of royal courts.","R","100+"),
    ("Renaissance","n.","文艺复兴","The Renaissance marked a rebirth of art and intellectual inquiry.","R L","90+"),
    ("migration","n.","迁移","Human migration shaped the genetic diversity of populations.","R L","80+"),
],
"艺术与文学 Art & Literature":[
    ("sculpture","n.","雕塑","Greek sculpture celebrated the idealized human form.","R L","80+"),
    ("narrative","n.","叙事","The novel employs a first-person narrative to engage readers.","R L","80+"),
    ("contemporary","adj.","当代的","Contemporary art challenges traditional aesthetic boundaries.","R L","80+"),
    ("portrait","n.","肖像","Renaissance portraits often conveyed the status of their subjects.","R L","80+"),
    ("genre","n.","体裁","Science fiction is a genre that explores futuristic ideas.","R L","90+"),
    ("aesthetic","adj.","审美的","The aesthetic principles of Japanese gardens emphasize harmony.","R","90+"),
    ("manuscript","n.","手稿","The medieval manuscript was decorated with elaborate illustrations.","R","90+"),
    ("masterpiece","n.","杰作","The painting is considered a masterpiece of Impressionist art.","R L","90+"),
    ("symbolism","n.","象征主义","Symbolism in literature uses objects to represent abstract ideas.","R","90+"),
    ("patron","n.","赞助人","Wealthy patrons funded many Renaissance artists and architects.","R","90+"),
    ("pottery","n.","陶器","Pottery fragments help archaeologists date ancient settlements.","R L","80+"),
    ("abstract","adj.","抽象的","Abstract art does not attempt to represent visual reality.","R L","90+"),
    ("literary","adj.","文学的","Literary analysis examines how authors use language and structure.","R","90+"),
    ("impressionism","n.","印象派","Impressionism captured fleeting effects of light and colour.","R","100+"),
    ("Renaissance","n.","文艺复兴","Renaissance artists pioneered techniques such as linear perspective.","R L","90+"),
],
"心理与行为 Psychology & Behavior":[
    ("cognitive","adj.","认知的","Cognitive psychology examines how people process information.","R L","80+"),
    ("motivation","n.","动机","Intrinsic motivation leads to deeper engagement with tasks.","R L","80+"),
    ("perception","n.","感知","Visual perception can be influenced by context and expectation.","R L","80+"),
    ("stimulus","n.","刺激","The brain processes each stimulus through specialized neural pathways.","R L","90+"),
    ("hypothesis","n.","假设","The researcher tested the hypothesis through controlled experiments.","R L","90+"),
    ("conditioning","n.","条件反射","Classical conditioning was first demonstrated by Pavlov.","R","90+"),
    ("consciousness","n.","意识","The nature of consciousness remains one of philosophy's deepest puzzles.","R","90+"),
    ("instinct","n.","本能","Newborns display certain survival instincts from birth.","R L","90+"),
    ("reinforcement","n.","强化","Positive reinforcement encourages the repetition of desired behaviour.","R","90+"),
    ("bias","n.","偏见","Confirmation bias leads people to favour information that supports their beliefs.","R L","90+"),
    ("neuron","n.","神经元","Neurons transmit electrical signals throughout the nervous system.","R","100+"),
    ("behavioral","adj.","行为的","Behavioral research tracks observable actions in controlled settings.","R L","80+"),
    ("developmental","adj.","发展的","Developmental psychology studies changes across the lifespan.","R","100+"),
    ("socialization","n.","社会化","Socialization shapes how individuals interact with their community.","R L","100+"),
    ("temperament","n.","气质","Research suggests temperament is partially determined by genetics.","R","100+"),
],
"社会与人类学 Society & Anthropology":[
    ("cultural","adj.","文化的","Cultural traditions are often passed down through oral storytelling.","R L","80+"),
    ("settlement","n.","定居点","The earliest permanent settlements emerged in fertile river valleys.","R L","80+"),
    ("urbanization","n.","城市化","Rapid urbanization creates significant infrastructure challenges.","R L","80+"),
    ("hierarchy","n.","等级制度","Social hierarchy in ancient societies was often determined by birth.","R L","80+"),
    ("ritual","n.","仪式","Ritual ceremonies serve important social and spiritual functions.","R L","90+"),
    ("anthropology","n.","人类学","Anthropology combines fieldwork and theory to understand cultures.","R","90+"),
    ("kinship","n.","亲属关系","Kinship systems define family obligations within a community.","R","90+"),
    ("subsistence","n.","生存","Early communities relied on subsistence farming for basic needs.","R L","90+"),
    ("nomadic","adj.","游牧的","Nomadic groups migrated seasonally to follow herds.","R L","90+"),
    ("assimilation","n.","同化","Assimilation policies suppressed minority languages and customs.","R L","90+"),
    ("demographic","adj.","人口统计的","Demographic shifts pose challenges for public policy.","R L","90+"),
    ("ethnography","n.","民族志","Ethnography involves immersive observation of communities.","R","100+"),
    ("stratification","n.","社会分层","Economic stratification leads to unequal access to education.","R","100+"),
    ("agrarian","adj.","农业的","The transition from an agrarian economy reshaped social structures.","R","100+"),
    ("egalitarian","adj.","平等主义的","Some hunter-gatherer societies maintained egalitarian structures.","R","100+"),
],
"环境与气候 Environment & Climate":[
    ("pollution","n.","污染","Industrial pollution has contaminated waterways worldwide.","R L","80+"),
    ("drought","n.","干旱","Prolonged drought led to crop failure across the region.","R L","80+"),
    ("conservation","n.","保护","Conservation efforts have helped restore endangered species.","R L","80+"),
    ("renewable","adj.","可再生的","Renewable energy sources such as wind and solar are expanding.","R L","80+"),
    ("ecosystem","n.","生态系统","Disrupting one component can cascade through the whole ecosystem.","R L","80+"),
    ("deforestation","n.","森林砍伐","Tropical deforestation contributes to biodiversity loss.","R L","90+"),
    ("greenhouse","adj.","温室的","Greenhouse gas concentrations have reached unprecedented levels.","R L","90+"),
    ("emission","n.","排放","Reducing carbon emissions is essential to slowing global warming.","R L","90+"),
    ("sustainability","n.","可持续性","Sustainability balances economic growth with environmental protection.","R L","90+"),
    ("biodiversity","n.","生物多样性","Biodiversity loss weakens ecosystems and reduces resilience.","R L","90+"),
    ("ozone","n.","臭氧","Ozone layer depletion increases exposure to ultraviolet radiation.","R","90+"),
    ("carbon","n.","碳","The carbon cycle exchanges carbon between atmosphere and land.","R L","90+"),
    ("desertification","n.","荒漠化","Desertification threatens agricultural land in arid regions.","R","100+"),
    ("reforestation","n.","重新造林","Reforestation programmes aim to absorb atmospheric CO2.","R","100+"),
    ("contamination","n.","污染","Groundwater contamination poses serious public health risks.","R","100+"),
],
"学术与教育 Academic & Education":[
    ("curriculum","n.","课程体系","The university revised its curriculum for interdisciplinary courses.","R L","80+"),
    ("seminar","n.","研讨会","Graduate students attend a weekly seminar on current research.","R L","80+"),
    ("enrollment","n.","注册","Online course enrollment has increased dramatically.","R L","80+"),
    ("scholarship","n.","奖学金","The scholarship covers tuition, housing, and living expenses.","R L","80+"),
    ("faculty","n.","教职员工","The faculty includes experts in biology, linguistics, and physics.","R L","80+"),
    ("syllabus","n.","教学大纲","The syllabus outlines readings, assignments, and grading criteria.","R L","90+"),
    ("dissertation","n.","学位论文","She spent three years on her doctoral dissertation.","R L","90+"),
    ("prerequisite","n.","先决条件","Calculus is a prerequisite for advanced physics.","R L","90+"),
    ("interdisciplinary","adj.","跨学科的","Interdisciplinary programmes combine insights from multiple fields.","R L","90+"),
    ("transcript","n.","成绩单","Applicants must submit an official transcript.","R L","90+"),
    ("plagiarism","n.","抄袭","Universities use software to detect plagiarism.","R L","90+"),
    ("accreditation","n.","认证","Accreditation ensures institutions meet quality standards.","R","100+"),
    ("pedagogy","n.","教学法","Research in pedagogy shows active learning improves outcomes.","R","100+"),
    ("tenure","n.","终身教职","Earning tenure requires a strong research record.","R L","100+"),
    ("symposium","n.","学术研讨会","The annual symposium brings scholars from around the world.","R","100+"),
],
"高频学术动词 High-Frequency Academic Verbs":[
    ("investigate","v.","调查","Scientists investigate the causes of colony collapse disorder.","R L","80+"),
    ("analyze","v.","分析","Researchers analyze data to identify behavioural patterns.","R L","80+"),
    ("demonstrate","v.","证明","The experiment demonstrated temperature affects reaction rates.","R L","80+"),
    ("illustrate","v.","说明","The case study illustrates how environment influences behaviour.","R L","80+"),
    ("enhance","v.","增强","New irrigation techniques enhanced crop yields significantly.","R L","80+"),
    ("derive","v.","源自","Many English words derive from Latin and Greek roots.","R","90+"),
    ("constitute","v.","构成","Small farms constitute the majority of agricultural operations.","R","90+"),
    ("perceive","v.","感知","Humans perceive colour through specialized receptor cells.","R L","90+"),
    ("facilitate","v.","促进","Trade routes facilitated exchange of goods and ideas.","R L","90+"),
    ("undermine","v.","破坏","Soil erosion can undermine the structural integrity of roads.","R L","90+"),
    ("advocate","v.","提倡","Conservation groups advocate for stricter emission regulations.","R L","90+"),
    ("evaluate","v.","评估","Peer reviewers evaluate submitted research methodology.","R L","90+"),
    ("correlate","v.","相关联","Regular exercise correlates with improved cognitive function.","R","100+"),
    ("synthesize","v.","合成","Plants synthesize glucose during photosynthesis.","R","100+"),
    ("replicate","v.","复制","Laboratories attempted to replicate the original results.","R","100+"),
],
}

def build():
    d=Document()
    sec=d.sections[0]; sec.page_width=Cm(29.7); sec.page_height=Cm(21)
    sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
    sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)

    for _ in range(3): s=d.add_paragraph();_sp(s,0,0)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"睿然国际教育",32,True,TEAL);_sp(p,0,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"RUIRAN INTERNATIONAL EDUCATION",11,False,GRAY,f="Arial");_sp(p,0,300)
    _hr(d,"2F8E87",8)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"TOEFL 托福核心词汇手册",28,True,TEAL_DARK);_sp(p,200,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"TPO 高频 · 学科精编 · 全科覆盖",16,False,TEAL);_sp(p,0,60)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"80+  ·  90+  ·  100+",13,True,GRAY,f="Arial");_sp(p,0,60)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"Reading  ·  Listening  ·  Speaking  ·  Writing",12,False,GRAY,f="Arial");_sp(p,0,200)
    _hr(d,"2F8E87",8)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"内部教学资料 · 请勿外传",10,False,ORANGE,i=True)
    d.add_page_break()

    # Usage
    hd(d,0,"使用说明","How to Use This Booklet")
    for t in ["本手册精选托福考试 TPO 高频核心词汇，按学科主题分类，覆盖阅读、听力、口语、写作四科。",
              "每个词条包含：单词/短语、词性、中文释义、TPO 风格例句、适用考试板块。"]:
        pp=d.add_paragraph();pp.paragraph_format.space_before=Pt(4);pp.paragraph_format.space_after=Pt(4)
        pp.paragraph_format.line_spacing=Pt(22);_r(pp,t,10.5,False,DARK)
    pp=d.add_paragraph();pp.paragraph_format.space_before=Pt(8)
    pPr=pp._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="2F8E87"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5F3" w:val="clear"/>'))
    _r(pp,"词汇按难度标注三个等级，以行背景色区分：\n",10,False,TEAL_DARK)
    _r(pp,"  ■ 80+ 基础词汇（暖黄底色）",10,True,B80);_r(pp,"   ",10)
    _r(pp,"■ 90+ 进阶词汇（青绿底色）",10,True,B90);_r(pp,"   ",10)
    _r(pp,"■ 100+ 高阶词汇（淡紫底色）",10,True,B100)
    pp2=d.add_paragraph();pp2.paragraph_format.space_before=Pt(8)
    _r(pp2,"考试板块缩写：",10,True,TEAL_DARK)
    _r(pp2," R = Reading　L = Listening　S = Speaking　W = Writing",10,False,GRAY,f="Arial")
    d.add_page_break()

    for idx,(topic,entries) in enumerate(V.items(),1):
        parts=topic.split(" ",1)
        hd(d,idx,parts[0],parts[1] if len(parts)>1 else "")
        vtable(d,entries)
        if idx<len(V): d.add_page_break()

    d.add_page_break()
    for _ in range(4): s=d.add_paragraph();_sp(s,0,0)
    _hr(d,"2F8E87",8)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"睿然国际教育",20,True,TEAL);_sp(p,160,80)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_r(p,"用科技赋能教育，让过程被看见，让结果自然发生。",11,False,GRAY,i=True)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;_sp(p,120,0)
    _r(p,"© 睿然国际教育  ·  内部教学资料  ·  版权所有",9,False,GRAY)
    _hr(d,"2F8E87",8)
    d.save(OUT);print(f"✅  Saved → {OUT}")

if __name__=="__main__": build()
